from __future__ import annotations

import re
import subprocess
import xml.etree.ElementTree as ET
from copy import deepcopy
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont, ImageOps
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_BREAK
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Pt
from docxcompose.composer import Composer
from pypdf import PdfReader, PdfWriter


PROJECT_ROOT = Path(__file__).resolve().parents[1]
DOC_OUTPUT_DIR = PROJECT_ROOT / "output" / "doc"
TMP_DOC_DIR = PROJECT_ROOT / "tmp" / "docs"
SCREENSHOT_DIR = PROJECT_ROOT / "docs" / "screenshots"

DOC_PATH = DOC_OUTPUT_DIR / "上海大学本科毕业论文_基于大数据的热门行业识别与龙头股遴选研究.docx"
TEMPLATE_PATH = Path("/Users/leonardodon/Downloads/上海大学本科毕业论文（设计）撰写格式模板.docx")
TEMPLATE_PDF_PATH = Path("/Users/leonardodon/Downloads/上海大学本科毕业论文（设计）撰写格式模板.pdf")
OUTPUT_PDF_PATH = DOC_OUTPUT_DIR / "上海大学本科毕业论文_基于大数据的热门行业识别与龙头股遴选研究.pdf"
LEGACY_DUPLICATE_PDF_PATH = DOC_OUTPUT_DIR / "上海大学本科毕业论文_基于大数据的热门行业识别与龙头股遴选研究_送审版.pdf"
COMPOSED_TMP_PATH = TMP_DOC_DIR / "shu_thesis_composed.docx"
TMP_TEMPLATE_PDF_DIR = TMP_DOC_DIR / "template_pdf_pages"
TMP_DOCX_RENDER_DIR = TMP_DOC_DIR / "current_docx_render"
TMP_SUBMISSION_RENDER_DIR = TMP_DOC_DIR / "submission_pdf_pages"
TMP_FRONT_ASSET_DIR = TMP_DOC_DIR / "pdf_front_assets"
TEMPLATE_XML_PREFIX = TMP_TEMPLATE_PDF_DIR / "template_layout"

THESIS_TITLE = "基于大数据的热门行业识别与龙头股遴选研究"
STUDENT_INFO = {
    "college": "通信与信息工程学院",
    "major": "通信工程",
    "student_id": "22121527",
    "student_name": "唐梓涵",
    "advisor": "沈文辉",
    "duration": "2025.12-2026.6",
}

FIGURE_IMAGES = {
    "图 5.1 行业热度总览界面": {
        "path": SCREENSHOT_DIR / "industry-ranking-overview.png",
        "crop": (250, 140, 1360, 1160),
        "slug": "figure_5_1",
        "width_scale": 0.72,
    },
    "图 5.2 行业热力图界面": {
        "path": SCREENSHOT_DIR / "industry-heatmap-overview.png",
        "crop": (240, 120, 1360, 1180),
        "slug": "figure_5_2",
    },
    "图 5.3 龙头股详情界面": {
        "path": SCREENSHOT_DIR / "leader-stock-detail.png",
        "crop": (300, 90, 1160, 1080),
        "slug": "figure_5_3",
    },
}
TMP_ASSET_DIR = TMP_DOC_DIR / "figure_assets"
ARCHITECTURE_FIGURE_PATH = TMP_ASSET_DIR / "figure_3_1_system_architecture.png"
CJK_FONT_SOURCES = (
    (Path("/Library/Fonts/SimSun.ttf"), 0),
    (Path("/System/Library/Fonts/SimSun.ttf"), 0),
    (Path("/System/Library/Fonts/Supplemental/Songti.ttc"), 6),
    (Path("/System/Library/Fonts/Supplemental/Songti.ttc"), 1),
    (Path("/System/Library/Fonts/Supplemental/Songti.ttc"), 4),
)
BODY_PDF_ABSTRACT_PAGE_INDEX = 4
BODY_FIRST_LINE_INDENT = Emu(266700)
BODY_LINE_SPACING = Pt(23)
ABSTRACT_LINE_SPACING = Pt(20)
SOURCE_NOTE = "图片来源：系统运行截图（作者自制）"
ARCHITECTURE_SOURCE_NOTE = "图片来源：作者根据系统实现绘制"
HEADER_TEXT = "上海大学本科毕业论文（设计）"

COVER_FIELD_ORDER = [
    "title",
    "college",
    "major",
    "student_id",
    "student_name",
    "advisor",
    "duration",
]

COVER_FIELD_SPECS = {
    "title": {
        "text": THESIS_TITLE,
        # SHU cover spec: title in small No.2, rendered here on the
        # 150dpi template image as an approximately equivalent pixel size.
        "font_size": 38,
        "max_width": 520,
        "align": "center",
        "padding_x": 0,
        "gap_above_line": 4,
    },
    "college": {
        "text": STUDENT_INFO["college"],
        "font_size": 33,
        "max_width": 438,
        "align": "left",
        "padding_x": 8,
        "gap_above_line": 2,
    },
    "major": {
        "text": STUDENT_INFO["major"],
        "font_size": 33,
        "max_width": 438,
        "align": "left",
        "padding_x": 8,
        "gap_above_line": 2,
    },
    "student_id": {
        "text": STUDENT_INFO["student_id"],
        "font_size": 33,
        "max_width": 440,
        "align": "left",
        "padding_x": 8,
        "gap_above_line": 2,
    },
    "student_name": {
        "text": STUDENT_INFO["student_name"],
        "font_size": 33,
        "max_width": 440,
        "align": "left",
        "padding_x": 8,
        "gap_above_line": 2,
    },
    "advisor": {
        "text": STUDENT_INFO["advisor"],
        "font_size": 33,
        "max_width": 440,
        "align": "left",
        "padding_x": 8,
        "gap_above_line": 2,
    },
    "duration": {
        "text": STUDENT_INFO["duration"],
        "font_size": 33,
        "max_width": 440,
        "align": "left",
        "padding_x": 8,
        "gap_above_line": 2,
    },
}

COVER_FALLBACK_LINES = {
    "title": {"x_start": 458, "x_end": 964, "y_line": 848},
    "college": {"x_start": 444, "x_end": 949, "y_line": 976},
    "major": {"x_start": 444, "x_end": 949, "y_line": 1041},
    "student_id": {"x_start": 444, "x_end": 949, "y_line": 1106},
    "student_name": {"x_start": 444, "x_end": 949, "y_line": 1171},
    "advisor": {"x_start": 444, "x_end": 949, "y_line": 1236},
    "duration": {"x_start": 444, "x_end": 949, "y_line": 1301},
}

DECLARATION_FIELD_SPECS = {
    "student_name": {
        "text": STUDENT_INFO["student_name"],
        "anchor_label": "姓名：",
        "font_size": 20,
        "max_width": 160,
        "padding_x": 18,
        "vertical_align": "center",
    },
    "student_id": {
        "text": STUDENT_INFO["student_id"],
        "anchor_label": "学号：",
        "font_size": 19,
        "max_width": 170,
        "padding_x": 16,
        "vertical_align": "center",
    },
    "thesis_title": {
        "text": THESIS_TITLE,
        "anchor_label": "论文题目：",
        "font_size": 20,
        "max_width": 560,
        "padding_x": 22,
        "vertical_align": "center",
    },
}

DECLARATION_FALLBACK_LABEL_BOXES = {
    "姓名：": {"left": 214, "top": 160, "right": 339, "bottom": 185},
    "学号：": {"left": 747, "top": 160, "right": 822, "bottom": 185},
    "论文题目：": {"left": 214, "top": 208, "right": 339, "bottom": 233},
}


TOC_BLUEPRINT = [
    ("toc 1", "摘  要"),
    ("toc 1", "ABSTRACT"),
    ("toc 1", "1 绪论"),
    ("toc 2", "1.1 研究背景与意义"),
    ("toc 2", "1.2 国内外研究现状"),
    ("toc 2", "1.3 研究内容与技术路线"),
    ("toc 2", "1.4 论文结构安排"),
    ("toc 1", "2 相关技术与理论基础"),
    ("toc 2", "2.1 金融大数据的特征"),
    ("toc 2", "2.2 多源数据采集与清洗"),
    ("toc 2", "2.3 热门行业识别的理论基础"),
    ("toc 2", "2.4 龙头股评价的理论基础"),
    ("toc 2", "2.5 系统关键技术"),
    ("toc 1", "3 系统需求分析与总体设计"),
    ("toc 2", "3.1 设计目标"),
    ("toc 2", "3.2 功能需求分析"),
    ("toc 2", "3.3 非功能需求分析"),
    ("toc 2", "3.4 系统总体架构设计"),
    ("toc 2", "3.5 数据流程设计"),
    ("toc 2", "3.6 存储设计"),
    ("toc 1", "4 热门行业识别与龙头股遴选模型设计"),
    ("toc 2", "4.1 数据预处理与指标标准化"),
    ("toc 2", "4.2 热门行业识别模型设计"),
    ("toc 2", "4.3 行业波动率估计与聚类辅助分析"),
    ("toc 2", "4.4 龙头股综合评分模型设计"),
    ("toc 2", "4.5 快速评分机制设计"),
    ("toc 1", "5 系统实现"),
    ("toc 2", "5.1 数据提供器与回退机制实现"),
    ("toc 2", "5.2 行业分析模块实现"),
    ("toc 2", "5.3 龙头股评分模块实现"),
    ("toc 2", "5.4 后端接口实现"),
    ("toc 2", "5.5 前端可视化实现"),
    ("toc 2", "5.6 偏好配置与持续跟踪实现"),
    ("toc 1", "6 系统测试与结果分析"),
    ("toc 2", "6.1 测试环境"),
    ("toc 2", "6.2 功能测试"),
    ("toc 2", "6.3 热力图快照结果分析"),
    ("toc 2", "6.4 系统特征与不足分析"),
    ("toc 2", "6.5 对毕业设计目标的达成情况"),
    ("toc 1", "结 论"),
    ("toc 1", "参考文献"),
    ("toc 1", "致 谢"),
]


def clear_paragraph(paragraph) -> None:
    p = paragraph._element
    p_pr = p.pPr
    for child in list(p):
        if child is not p_pr:
            p.remove(child)


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def remove_section_break(paragraph) -> None:
    p_pr = paragraph._element.pPr
    if p_pr is None:
        return
    sect_pr = p_pr.find(qn("w:sectPr"))
    if sect_pr is not None:
        p_pr.remove(sect_pr)


def set_east_asia_font(run, font_name: str) -> None:
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)
    r_fonts.set(qn("w:ascii"), "Times New Roman")
    r_fonts.set(qn("w:hAnsi"), "Times New Roman")


def replace_paragraph_text(paragraph, text: str) -> None:
    clear_paragraph(paragraph)
    paragraph.add_run(text)


def replace_cover_line(paragraph, prefix: str, value: str, gap_spaces: int = 8) -> None:
    clear_paragraph(paragraph)
    paragraph.add_run(f"{prefix}{' ' * gap_spaces}{value}")


def ensure_page_break_before(paragraph) -> None:
    previous = paragraph._element.getprevious()
    if previous is not None and previous.tag == qn("w:p"):
        from docx.text.paragraph import Paragraph

        previous_para = Paragraph(previous, paragraph._parent)
        if previous_para.text.strip() == "" and 'w:type="page"' in previous_para._element.xml:
            return
    break_para = paragraph.insert_paragraph_before()
    break_para.add_run().add_break(WD_BREAK.PAGE)


def remove_page_break_before_flag(paragraph) -> None:
    p_pr = paragraph._element.get_or_add_pPr()
    for node in p_pr.findall(qn("w:pageBreakBefore")):
        p_pr.remove(node)


def format_unnumbered_heading(paragraph) -> None:
    paragraph.style = "Normal"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.0
    if not paragraph.runs:
        paragraph.add_run(paragraph.text)
    for run in paragraph.runs:
        run.font.bold = True
        run.font.size = Pt(18)
        run.font.name = "Times New Roman"
        set_east_asia_font(run, "黑体")


def format_toc_title(paragraph) -> None:
    clear_paragraph(paragraph)
    paragraph.style = "Normal"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = Pt(20)
    paragraph.paragraph_format.page_break_before = True
    run = paragraph.add_run("目  录")
    run.font.bold = False
    run.font.size = Pt(18)
    run.font.name = "Times New Roman"
    set_east_asia_font(run, "黑体")


def add_page_number_field(paragraph) -> None:
    clear_paragraph(paragraph)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run = paragraph.add_run()
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)
    run._r.append(fld_end)
    run.font.name = "Times New Roman"
    run.font.size = Pt(10.5)
    set_east_asia_font(run, "宋体")


def set_section_page_number_format(section, fmt: str | None = None, start: int | None = None) -> None:
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    if fmt:
        pg_num.set(qn("w:fmt"), fmt)
    if start is not None:
        pg_num.set(qn("w:start"), str(start))


def configure_section_header(section, text: str | None) -> None:
    section.different_first_page_header_footer = False
    header_para = section.header.paragraphs[0]
    clear_paragraph(header_para)
    header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_para.paragraph_format.first_line_indent = Pt(0)
    header_para.paragraph_format.space_before = Pt(0)
    header_para.paragraph_format.space_after = Pt(0)
    if text:
        run = header_para.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(10.5)
        run.font.bold = False
        set_east_asia_font(run, "宋体")


def configure_section_footer(section, with_page_number: bool) -> None:
    footer_para = section.footer.paragraphs[0]
    clear_paragraph(footer_para)
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.paragraph_format.first_line_indent = Pt(0)
    footer_para.paragraph_format.space_before = Pt(0)
    footer_para.paragraph_format.space_after = Pt(0)
    if with_page_number:
        add_page_number_field(footer_para)


def set_run_text_font(run, east_asia_font: str, size: Pt, bold: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run.font.size = size
    if bold is not None:
        run.font.bold = bold
    set_east_asia_font(run, east_asia_font)


def format_body_run(run) -> None:
    if not run.text:
        return
    if run.font.size is None:
        run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    set_east_asia_font(run, "宋体")


def normalize_search_text(text: str) -> str:
    normalized = re.sub(r"\s+", "", text or "")
    return normalized.replace("（", "(").replace("）", ")")


def insert_paragraph_after(paragraph, text: str = ""):
    from docx.text.paragraph import Paragraph

    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    return new_para


def insert_paragraph_before_element(element, parent):
    from docx.text.paragraph import Paragraph

    new_p = OxmlElement("w:p")
    element.addprevious(new_p)
    return Paragraph(new_p, parent)


def fill_cover(doc: Document) -> None:
    replacements = {
        "题   目：": (THESIS_TITLE, 8),
        "学    院：": (STUDENT_INFO["college"], 8),
        "专    业：": (STUDENT_INFO["major"], 8),
        "学    号：": (STUDENT_INFO["student_id"], 8),
        "学生姓名：": (STUDENT_INFO["student_name"], 6),
        "指导教师：": (STUDENT_INFO["advisor"], 6),
        "起讫日期：": (STUDENT_INFO["duration"], 8),
    }
    for paragraph in doc.paragraphs:
        for prefix, (value, gap_spaces) in replacements.items():
            if paragraph.text.startswith(prefix):
                replace_cover_line(paragraph, prefix, value, gap_spaces=gap_spaces)


def fill_declaration_table(doc: Document) -> None:
    if not doc.tables:
        return
    table = doc.tables[0]
    table.cell(0, 0).text = f"姓    名：{STUDENT_INFO['student_name']}"
    table.cell(0, 1).text = f"学号：{STUDENT_INFO['student_id']}"
    merged = table.cell(1, 0).merge(table.cell(1, 1))
    merged.text = f"论文题目：{THESIS_TITLE}"


def remove_elements_before_paragraph(doc: Document, text: str) -> None:
    body = doc._element.body
    anchor = None
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == text:
            anchor = paragraph._element
            break
    if anchor is None:
        raise RuntimeError(f"Paragraph not found for strip-before: {text}")

    for child in list(body):
        if child is anchor:
            break
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def remove_elements_from_paragraph(doc: Document, text: str) -> None:
    body = doc._element.body
    anchor = None
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == text:
            anchor = paragraph._element
            break
    if anchor is None:
        raise RuntimeError(f"Paragraph not found for strip-after: {text}")

    removing = False
    for child in list(body):
        if child is anchor:
            removing = True
        if removing and child.tag != qn("w:sectPr"):
            body.remove(child)


def find_heading_break_paragraph(doc: Document, heading_text: str):
    heading_index = None
    for index, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == heading_text:
            heading_index = index
            break
    if heading_index is None:
        raise RuntimeError(f"Heading not found: {heading_text}")

    for index in range(heading_index - 1, max(-1, heading_index - 5), -1):
        if "w:sectPr" in doc.paragraphs[index]._element.xml:
            return doc.paragraphs[index]
    raise RuntimeError(f"Section break paragraph not found before heading: {heading_text}")


def replace_first_paragraph_starting_with(doc: Document, prefix: str, replacement: str) -> bool:
    prefix_hint = prefix[:24]
    replacement_prefix = replacement[:24]
    for paragraph in doc.paragraphs:
        stripped = paragraph.text.strip()
        if stripped.startswith(prefix) or stripped.startswith(prefix_hint) or stripped.startswith(replacement_prefix):
            replace_paragraph_text(paragraph, replacement)
            return True
    return False


def find_last_paragraph_index_by_text(doc: Document, text: str) -> int:
    matches = [idx for idx, paragraph in enumerate(doc.paragraphs) if paragraph.text.strip() == text]
    if not matches:
        raise RuntimeError(f"Paragraph not found: {text}")
    return matches[-1]


def replace_section_body(doc: Document, heading_text: str, next_heading_text: str, paragraphs: list[str]) -> None:
    start_index = find_last_paragraph_index_by_text(doc, heading_text)
    end_index = next(
        (
            idx
            for idx in range(start_index + 1, len(doc.paragraphs))
            if doc.paragraphs[idx].text.strip() == next_heading_text
        ),
        None,
    )
    if end_index is None:
        raise RuntimeError(f"Section end heading not found after {heading_text}: {next_heading_text}")

    heading_para = doc.paragraphs[start_index]
    end_element = doc.paragraphs[end_index]._element
    current_element = heading_para._element.getnext()
    while current_element is not None and current_element is not end_element:
        next_element = current_element.getnext()
        current_element.getparent().remove(current_element)
        current_element = next_element

    current = heading_para
    for text in paragraphs:
        current = insert_paragraph_after(current, text)
        if re.match(r"^\d+\.\d+\.\d+\s+", text):
            current.style = "Heading 3"


def insert_table_after(doc: Document, paragraph, rows: int, cols: int):
    table = doc.add_table(rows=rows, cols=cols)
    tbl = table._tbl
    paragraph._element.addnext(tbl)
    return table


def set_cell_fill(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color: str | None = None, size: str = "16") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = tc_borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_borders.append(node)
        if color:
            node.set(qn("w:val"), "single")
            node.set(qn("w:sz"), size)
            node.set(qn("w:space"), "0")
            node.set(qn("w:color"), color)
        else:
            node.set(qn("w:val"), "nil")


def format_architecture_cell(cell, title: str, body_lines: list[str], fill: str, border_color: str) -> None:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_fill(cell, fill)
    set_cell_borders(cell, border_color, size="18")
    paragraph = cell.paragraphs[0]
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = Pt(17)
    title_run = paragraph.add_run(title)
    set_run_text_font(title_run, "宋体", Pt(16), bold=True)
    title_run.add_break(WD_BREAK.LINE)
    body_run = paragraph.add_run("\n".join(body_lines))
    set_run_text_font(body_run, "宋体", Pt(12))


def format_architecture_arrow_cell(cell) -> None:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_fill(cell, "F8FAFD")
    set_cell_borders(cell, None)
    paragraph = cell.paragraphs[0]
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("↓")
    set_run_text_font(run, "宋体", Pt(16), bold=True)


def format_architecture_summary_cell(cell, lines: list[str]) -> None:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_fill(cell, "FFFFFF")
    set_cell_borders(cell, "B0B8C2", size="12")
    paragraph = cell.paragraphs[0]
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = Pt(16)
    run = paragraph.add_run("\n".join(lines))
    set_run_text_font(run, "宋体", Pt(11))


def build_architecture_diagram(output_path: Path) -> Path:
    TMP_ASSET_DIR.mkdir(parents=True, exist_ok=True)
    width, height = 1280, 980
    image = Image.new("RGB", (width, height), "#F8FAFD")
    draw = ImageDraw.Draw(image)
    font_sources = get_available_cjk_font_sources()

    title_font = fit_cjk_font(draw, "系统总体架构图", 520, 34, min_size=28, font_sources=font_sources)
    layer_font = fit_cjk_font(draw, "表现层", 260, 28, min_size=22, font_sources=font_sources)
    text_font = fit_cjk_font(draw, "行业热力图 / 行业排行榜", 920, 22, min_size=18, font_sources=font_sources)
    summary_font = fit_cjk_font(draw, "多源适配、分层缓存与前后端协同", 900, 20, min_size=16, font_sources=font_sources)

    draw.text((width / 2, 44), "系统总体架构图", fill="#243447", font=title_font, anchor="ma")

    layer_specs = [
        ("表现层", "行业热力图 / 行业排行榜\n行业趋势面板 / 龙头股详情", "#D9EAFB", "#5B8FD9"),
        ("服务层", "FastAPI 行业接口 / 参数校验\n响应封装 / 端点缓存", "#E3F4E8", "#5BA36B"),
        ("分析层", "行业分析器 / 龙头股评分器\n快速评分链路 / 波动率估计", "#FDEBCF", "#C58B1C"),
        ("数据层", "THS 主数据 / AKShare 增强\nSina 与腾讯补充 / JSON 与 localStorage", "#EFE2FB", "#8B62C4"),
    ]

    box_left, box_right = 120, 1160
    box_height = 145
    gap = 28
    start_top = 88

    for index, (layer_name, layer_text, fill_color, line_color) in enumerate(layer_specs):
        top = start_top + index * (box_height + gap)
        bottom = top + box_height
        draw.rounded_rectangle((box_left, top, box_right, bottom), radius=28, fill=fill_color, outline=line_color, width=5)
        draw.text((width / 2, top + 42), layer_name, fill="#1F2D3D", font=layer_font, anchor="ma")
        draw.multiline_text(
            (width / 2, top + 88),
            layer_text,
            fill="#2E3A46",
            font=text_font,
            anchor="ma",
            align="center",
            spacing=10,
        )
        if index < len(layer_specs) - 1:
            arrow_center_x = width // 2
            arrow_top = bottom + 8
            arrow_bottom = bottom + gap - 4
            draw.line((arrow_center_x, arrow_top, arrow_center_x, arrow_bottom), fill="#7A8794", width=5)
            draw.polygon(
                [
                    (arrow_center_x, arrow_bottom + 10),
                    (arrow_center_x - 12, arrow_bottom - 6),
                    (arrow_center_x + 12, arrow_bottom - 6),
                ],
                fill="#7A8794",
            )

    draw.rounded_rectangle((185, 865, 1095, 935), radius=20, fill="#FFFFFF", outline="#B0B8C2", width=3)
    draw.multiline_text(
        (width / 2, 898),
        "多源适配、分层缓存与前后端协同\n共同支撑热门行业识别与龙头股遴选闭环",
        fill="#3A4652",
        font=summary_font,
        anchor="mm",
        align="center",
        spacing=10,
    )

    image.save(output_path)
    return output_path


def insert_architecture_figure(doc: Document) -> None:
    anchor = find_last_paragraph_by_text(doc, "3.5 数据流程设计")
    break_para = insert_paragraph_before_element(anchor._element, anchor._parent)
    break_para.paragraph_format.first_line_indent = Pt(0)
    break_para.paragraph_format.space_before = Pt(0)
    break_para.paragraph_format.space_after = Pt(0)
    break_para.paragraph_format.page_break_before = True
    break_para.paragraph_format.keep_with_next = True

    diagram_table = insert_table_after(doc, break_para, rows=8, cols=1)
    diagram_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    diagram_table.autofit = False
    table_width = int((doc.sections[-1].page_width - doc.sections[-1].left_margin - doc.sections[-1].right_margin) * 0.88)
    for row in diagram_table.rows:
        row.cells[0].width = Emu(table_width)
        _set_row_no_split(row)

    row_specs = [
        ("layer", "表现层", ["行业热力图 / 行业排行榜", "行业趋势面板 / 龙头股详情"], "D9EAFB", "5B8FD9", Cm(2.2)),
        ("arrow",),
        ("layer", "服务层", ["FastAPI 行业接口 / 参数校验", "响应封装 / 端点缓存"], "E3F4E8", "5BA36B", Cm(2.1)),
        ("arrow",),
        ("layer", "分析层", ["行业分析器 / 龙头股评分器", "快速评分 / 波动率估计"], "FDEBCF", "C58B1C", Cm(2.1)),
        ("arrow",),
        ("layer", "数据层", ["THS 主数据 / AKShare 增强", "Sina 与腾讯补充 / JSON 与 localStorage"], "EFE2FB", "8B62C4", Cm(2.2)),
        ("summary", ["多源适配、分层缓存与前后端协同", "共同支撑热门行业识别与龙头股遴选闭环"], Cm(1.55)),
    ]

    for row, spec in zip(diagram_table.rows, row_specs):
        kind = spec[0]
        if kind == "layer":
            _, title, lines, fill, border, height = spec
            row.height = height
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            format_architecture_cell(row.cells[0], title, lines, fill, border)
        elif kind == "summary":
            _, lines, height = spec
            row.height = height
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            format_architecture_summary_cell(row.cells[0], lines)
        else:
            row.height = Cm(0.45)
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            format_architecture_arrow_cell(row.cells[0])

    caption = insert_paragraph_before_element(anchor._element, anchor._parent)
    caption.add_run("图 3.1 系统总体架构图")
    caption.paragraph_format.keep_with_next = True
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.first_line_indent = Pt(0)
    caption.paragraph_format.space_before = Pt(6)
    caption.paragraph_format.space_after = Pt(0)
    for run in caption.runs:
        set_run_text_font(run, "宋体", Pt(12))

    source = insert_paragraph_before_element(anchor._element, anchor._parent)
    source.add_run(ARCHITECTURE_SOURCE_NOTE)
    source.alignment = WD_ALIGN_PARAGRAPH.CENTER
    source.paragraph_format.first_line_indent = Pt(0)
    source.paragraph_format.space_before = Pt(0)
    source.paragraph_format.space_after = Pt(0)
    source.paragraph_format.keep_with_next = True
    for run in source.runs:
        set_run_text_font(run, "宋体", Pt(12), bold=True)


def insert_task_completion_table(doc: Document) -> None:
    heading = find_last_paragraph_by_text(doc, "6.5 对毕业设计目标的达成情况")
    intro_para = None
    next_element = heading._element.getnext()
    if next_element is not None and next_element.tag == qn("w:p"):
        from docx.text.paragraph import Paragraph

        intro_para = Paragraph(next_element, heading._parent)
    if intro_para is None:
        intro_para = insert_paragraph_after(heading, "对照任务书要求，如表 6.3 所示，本文在理论分析、模型设计和系统实现三个层面均完成了核心目标。")

    caption = insert_paragraph_after(intro_para, "表 6.3 毕业设计任务书目标达成情况")
    caption.paragraph_format.keep_with_next = True

    table = insert_table_after(doc, caption, rows=4, cols=4)
    rows = [
        ("任务书要求", "论文对应内容", "系统对应实现", "达成情况"),
        ("理解金融大数据挖掘相关技术与研究现状", "第 1 章与第 2 章完成文献综述和理论基础梳理", "明确行业子系统的数据来源、分析流程与关键技术", "已达成"),
        ("掌握行业识别与龙头股筛选方法", "第 4 章构建行业热度模型与龙头股评分模型", "实现行业横截面评分、波动率代理与双路径个股评分", "已达成"),
        ("编程设计软件系统跟踪重点行业和关键企业", "第 5 章与第 6 章说明系统实现与验证过程", "完成 THS-first、FastAPI、React、热力图快照与偏好持久化闭环", "已达成"),
    ]
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            table.cell(r_idx, c_idx).text = value


def ensure_blank_page_before_declaration(doc: Document) -> None:
    if not doc.tables:
        return

    declaration_table = doc.tables[0]

    next_sibling = declaration_table._element.getnext()
    if next_sibling is not None and next_sibling.tag == qn("w:p") and 'w:type="page"' in next_sibling.xml:
        next_sibling.getparent().remove(next_sibling)

    previous_sibling = declaration_table._element.getprevious()
    if previous_sibling is not None and previous_sibling.tag == qn("w:p") and 'w:type="page"' in previous_sibling.xml:
        return

    break_para = OxmlElement("w:p")
    run = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run.append(br)
    break_para.append(run)
    declaration_table._element.addprevious(break_para)


def set_core_properties(doc: Document) -> None:
    doc.core_properties.author = STUDENT_INFO["student_name"]
    doc.core_properties.title = THESIS_TITLE
    doc.core_properties.subject = THESIS_TITLE
    doc.core_properties.keywords = "金融大数据, 热门行业识别, 龙头股遴选, 量化研究平台"


def run_checked(command: list[str]) -> None:
    subprocess.run(command, check=True)


def get_available_cjk_font_sources() -> list[tuple[Path, int]]:
    sources = [(path, index) for path, index in CJK_FONT_SOURCES if path.exists()]
    if not sources:
        raise RuntimeError("No available CJK font found for thesis front-page rendering.")
    return sources


def load_cjk_font(size: int, font_sources: list[tuple[Path, int]] | None = None):
    sources = font_sources or get_available_cjk_font_sources()
    for path, index in sources:
        try:
            return ImageFont.truetype(str(path), size=size, index=index)
        except OSError:
            continue
    raise RuntimeError("Failed to load a usable CJK font source for thesis rendering.")


def fit_cjk_font(
    draw: ImageDraw.ImageDraw,
    text: str,
    max_width: int,
    start_size: int,
    min_size: int = 14,
    font_sources: list[tuple[Path, int]] | None = None,
):
    for size in range(start_size, min_size - 1, -1):
        font = load_cjk_font(size, font_sources=font_sources)
        bbox = draw.textbbox((0, 0), text, font=font)
        if bbox[2] - bbox[0] <= max_width:
            return font
    return load_cjk_font(min_size, font_sources=font_sources)


def make_box(left: float, top: float, right: float, bottom: float) -> dict[str, float]:
    return {
        "left": float(left),
        "top": float(top),
        "right": float(right),
        "bottom": float(bottom),
    }


def merge_boxes(boxes: list[dict[str, float]]) -> dict[str, float]:
    return make_box(
        min(box["left"] for box in boxes),
        min(box["top"] for box in boxes),
        max(box["right"] for box in boxes),
        max(box["bottom"] for box in boxes),
    )


def scale_box(box: dict[str, float], scale_x: float, scale_y: float) -> dict[str, float]:
    return make_box(
        box["left"] * scale_x,
        box["top"] * scale_y,
        box["right"] * scale_x,
        box["bottom"] * scale_y,
    )


def normalize_pdf_text(text: str) -> str:
    return re.sub(r"\s+", "", text or "")


def ensure_template_pdf_xml() -> Path:
    TMP_TEMPLATE_PDF_DIR.mkdir(parents=True, exist_ok=True)
    existing = sorted(TMP_TEMPLATE_PDF_DIR.glob("template_layout*.xml"))
    if existing:
        return existing[0]

    run_checked([
        "pdftohtml",
        "-xml",
        "-f",
        "1",
        "-l",
        "2",
        "-nomerge",
        "-noroundcoord",
        str(TEMPLATE_PDF_PATH),
        str(TEMPLATE_XML_PREFIX),
    ])
    generated = sorted(TMP_TEMPLATE_PDF_DIR.glob("template_layout*.xml"))
    if not generated:
        raise RuntimeError("Failed to extract template XML layout for thesis front pages.")
    return generated[0]


def ensure_template_pdf_pages() -> tuple[Path, Path]:
    TMP_TEMPLATE_PDF_DIR.mkdir(parents=True, exist_ok=True)
    page_1 = TMP_TEMPLATE_PDF_DIR / "template_page-01.png"
    page_2 = TMP_TEMPLATE_PDF_DIR / "template_page-02.png"
    if page_1.exists() and page_2.exists():
        return page_1, page_2

    prefix = TMP_TEMPLATE_PDF_DIR / "template_page"
    run_checked([
        "pdftoppm",
        "-f",
        "1",
        "-l",
        "2",
        "-png",
        str(TEMPLATE_PDF_PATH),
        str(prefix),
    ])
    if not page_1.exists() or not page_2.exists():
        raise RuntimeError("Failed to render official template PDF pages.")
    return page_1, page_2


def detect_cover_lines_in_image(image: Image.Image) -> dict[str, dict[str, int]]:
    grayscale = image.convert("L")
    width, height = grayscale.size
    search_left = int(width * 0.34)
    search_right = int(width * 0.80)
    search_top = int(height * 0.43)
    search_bottom = int(height * 0.75)
    min_length = int(width * 0.20)
    candidates: list[tuple[int, int, int]] = []

    for y in range(search_top, search_bottom):
        best_start = None
        best_end = None
        run_start = None
        for x in range(search_left, search_right + 1):
            is_dark = grayscale.getpixel((x, y)) < 120
            if is_dark:
                if run_start is None:
                    run_start = x
            elif run_start is not None:
                if best_start is None or x - run_start > best_end - best_start:
                    best_start, best_end = run_start, x - 1
                run_start = None
        if run_start is not None:
            if best_start is None or search_right - run_start > best_end - best_start:
                best_start, best_end = run_start, search_right
        if best_start is not None and best_end is not None and best_end - best_start >= min_length:
            candidates.append((y, best_start, best_end))

    groups: list[list[tuple[int, int, int]]] = []
    for row in candidates:
        if not groups or row[0] - groups[-1][-1][0] > 3:
            groups.append([row])
        else:
            groups[-1].append(row)

    normalized_groups = []
    for group in groups:
        line = {
            "y_line": int(round(sum(item[0] for item in group) / len(group))),
            "x_start": int(round(sum(item[1] for item in group) / len(group))),
            "x_end": int(round(sum(item[2] for item in group) / len(group))),
        }
        if line["x_end"] - line["x_start"] >= min_length:
            normalized_groups.append(line)

    if len(normalized_groups) < len(COVER_FIELD_ORDER):
        raise RuntimeError("Failed to detect all cover underline anchors from the official template.")
    if len(normalized_groups) > len(COVER_FIELD_ORDER):
        normalized_groups = sorted(
            normalized_groups,
            key=lambda line: line["x_end"] - line["x_start"],
            reverse=True,
        )[:len(COVER_FIELD_ORDER)]
        normalized_groups.sort(key=lambda line: line["y_line"])

    return {
        key: line
        for key, line in zip(COVER_FIELD_ORDER, normalized_groups[: len(COVER_FIELD_ORDER)])
    }


def detect_cover_lines(template_page: Path) -> dict[str, dict[str, int]]:
    try:
        with Image.open(template_page) as image:
            return detect_cover_lines_in_image(image)
    except Exception:
        return deepcopy(COVER_FALLBACK_LINES)


def extract_template_text_boxes(
    template_xml_path: Path,
    template_page: Path,
    page_number: int,
    labels: list[str],
) -> dict[str, dict[str, float]]:
    tree = ET.parse(template_xml_path)
    root = tree.getroot()
    page = next(
        (node for node in root.findall("page") if node.get("number") == str(page_number)),
        None,
    )
    if page is None:
        raise RuntimeError(f"Template XML page {page_number} not found.")

    with Image.open(template_page) as page_image:
        width, height = page_image.size
    scale_x = width / float(page.get("width"))
    scale_y = height / float(page.get("height"))

    nodes = []
    for element in page.findall("text"):
        raw_text = "".join(element.itertext())
        normalized = normalize_pdf_text(raw_text)
        if not normalized:
            continue
        left = float(element.get("left"))
        top = float(element.get("top"))
        box = make_box(
            left * scale_x,
            top * scale_y,
            (left + float(element.get("width"))) * scale_x,
            (top + float(element.get("height"))) * scale_y,
        )
        nodes.append({
            "text": raw_text,
            "normalized": normalized,
            "left_raw": left,
            "top_raw": top,
            "box": box,
        })

    lines: list[list[dict[str, object]]] = []
    for node in sorted(nodes, key=lambda item: (item["top_raw"], item["left_raw"])):
        if not lines or abs(node["top_raw"] - lines[-1][0]["top_raw"]) > 4:
            lines.append([node])
        else:
            lines[-1].append(node)

    boxes: dict[str, dict[str, float]] = {}
    for label in labels:
        normalized_label = normalize_pdf_text(label)
        for line in lines:
            sorted_line = sorted(line, key=lambda item: item["left_raw"])
            for start in range(len(sorted_line)):
                matched_nodes = []
                candidate = ""
                for end in range(start, len(sorted_line)):
                    piece = sorted_line[end]["normalized"]
                    next_candidate = candidate + piece
                    if not normalized_label.startswith(next_candidate):
                        break
                    candidate = next_candidate
                    matched_nodes.append(sorted_line[end]["box"])
                    if candidate == normalized_label:
                        boxes[label] = merge_boxes(matched_nodes)
                        break
                if label in boxes:
                    break
            if label in boxes:
                break

    for label in labels:
        if label not in boxes and label in DECLARATION_FALLBACK_LABEL_BOXES:
            boxes[label] = deepcopy(DECLARATION_FALLBACK_LABEL_BOXES[label])
    missing = [label for label in labels if label not in boxes]
    if missing:
        raise RuntimeError(f"Failed to extract template label boxes: {', '.join(missing)}")
    return boxes


def draw_text_above_line(
    draw: ImageDraw.ImageDraw,
    text: str,
    line_box: dict[str, int],
    spec: dict[str, object],
    font_sources: list[tuple[Path, int]],
) -> dict[str, float]:
    line_width = line_box["x_end"] - line_box["x_start"]
    padding = int(spec["padding_x"])
    available_width = min(int(spec["max_width"]), line_width - 2 * padding)
    font = fit_cjk_font(
        draw,
        text,
        available_width,
        int(spec["font_size"]),
        font_sources=font_sources,
    )
    bbox = draw.textbbox((0, 0), text, font=font)
    text_width = bbox[2] - bbox[0]
    if spec["align"] == "center":
        desired_left = line_box["x_start"] + (line_width - text_width) / 2
    else:
        desired_left = line_box["x_start"] + padding
    desired_bottom = line_box["y_line"] - int(spec["gap_above_line"])
    draw_x = desired_left - bbox[0]
    draw_y = desired_bottom - bbox[3]
    draw.text((draw_x, draw_y), text, font=font, fill=(18, 18, 18, 255))
    return make_box(
        draw_x + bbox[0],
        draw_y + bbox[1],
        draw_x + bbox[2],
        draw_y + bbox[3],
    )


def draw_text_after_label(
    draw: ImageDraw.ImageDraw,
    text: str,
    label_box: dict[str, float],
    spec: dict[str, object],
    font_sources: list[tuple[Path, int]],
) -> dict[str, float]:
    font = fit_cjk_font(
        draw,
        text,
        int(spec["max_width"]),
        int(spec["font_size"]),
        font_sources=font_sources,
    )
    bbox = draw.textbbox((0, 0), text, font=font)
    desired_left = label_box["right"] + int(spec["padding_x"])
    label_center_y = (label_box["top"] + label_box["bottom"]) / 2
    desired_center_y = label_center_y
    draw_x = desired_left - bbox[0]
    draw_y = desired_center_y - (bbox[1] + bbox[3]) / 2
    draw.text((draw_x, draw_y), text, font=font, fill=(18, 18, 18, 255))
    return make_box(
        draw_x + bbox[0],
        draw_y + bbox[1],
        draw_x + bbox[2],
        draw_y + bbox[3],
    )


def render_cover_page(template_page: Path, output_path: Path) -> dict[str, object]:
    line_boxes = detect_cover_lines(template_page)
    font_sources = get_available_cjk_font_sources()
    image = Image.open(template_page).convert("RGBA")
    draw = ImageDraw.Draw(image)
    placements = {}
    for key in COVER_FIELD_ORDER:
        placements[key] = draw_text_above_line(
            draw,
            str(COVER_FIELD_SPECS[key]["text"]),
            line_boxes[key],
            COVER_FIELD_SPECS[key],
            font_sources,
        )
    output_path.parent.mkdir(parents=True, exist_ok=True)
    image.save(output_path)
    return {
        "placements": placements,
        "anchors": line_boxes,
        "image_size": image.size,
    }


def render_declaration_page(template_page: Path, output_path: Path) -> dict[str, object]:
    font_sources = get_available_cjk_font_sources()
    template_xml = ensure_template_pdf_xml()
    label_boxes = extract_template_text_boxes(
        template_xml,
        template_page,
        2,
        [spec["anchor_label"] for spec in DECLARATION_FIELD_SPECS.values()],
    )
    image = Image.open(template_page).convert("RGBA")
    draw = ImageDraw.Draw(image)
    placements = {}
    for key, spec in DECLARATION_FIELD_SPECS.items():
        placements[key] = draw_text_after_label(
            draw,
            str(spec["text"]),
            label_boxes[str(spec["anchor_label"])],
            spec,
            font_sources,
        )
    output_path.parent.mkdir(parents=True, exist_ok=True)
    image.save(output_path)
    return {
        "placements": placements,
        "label_boxes": label_boxes,
        "image_size": image.size,
    }


def build_front_pdf(front_cover_png: Path, declaration_png: Path, output_pdf: Path) -> None:
    first = Image.open(front_cover_png).convert("RGB")
    second = Image.open(declaration_png).convert("RGB")
    output_pdf.parent.mkdir(parents=True, exist_ok=True)
    first.save(output_pdf, "PDF", save_all=True, append_images=[second], resolution=150.0)


def export_docx_pdf(doc_path: Path, output_dir: Path) -> Path:
    output_dir.mkdir(parents=True, exist_ok=True)
    run_checked([
        "soffice",
        "-env:UserInstallation=file:///tmp/lo_profile_shu_thesis",
        "--headless",
        "--convert-to",
        "pdf",
        "--outdir",
        str(output_dir),
        str(doc_path),
    ])
    pdf_path = output_dir / f"{doc_path.stem}.pdf"
    if not pdf_path.exists():
        raise RuntimeError(f"Failed to export PDF from DOCX: {pdf_path}")
    return pdf_path


def merge_submission_pdf(front_pdf: Path, body_pdf: Path, output_pdf: Path) -> None:
    front_reader = PdfReader(str(front_pdf))
    body_reader = PdfReader(str(body_pdf))
    writer = PdfWriter()

    for page in front_reader.pages:
        writer.add_page(page)
    for page in body_reader.pages[BODY_PDF_ABSTRACT_PAGE_INDEX:]:
        writer.add_page(page)

    with output_pdf.open("wb") as fp:
        writer.write(fp)


def render_pdf_preview(input_pdf: Path, output_dir: Path) -> None:
    output_dir.mkdir(parents=True, exist_ok=True)
    prefix = output_dir / "page"
    run_checked([
        "pdftoppm",
        "-f",
        "1",
        "-l",
        "40",
        "-png",
        str(input_pdf),
        str(prefix),
    ])


def get_rendered_preview_page(output_dir: Path, page_number: int) -> Path:
    candidates = [
        output_dir / f"page-{page_number:02d}.png",
        output_dir / f"page-{page_number}.png",
    ]
    for candidate in candidates:
        if candidate.exists():
            return candidate
    matches = sorted(output_dir.glob(f"page-*{page_number:02d}.png"))
    if matches:
        return matches[0]
    raise RuntimeError(f"Rendered preview page not found: {page_number}")


def validate_front_page_alignment(
    rendered_page: Path,
    placements: dict[str, dict[str, float]],
    source_size: tuple[int, int],
) -> None:
    with Image.open(rendered_page) as image:
        rendered_lines = detect_cover_lines_in_image(image)
        scale_x = image.size[0] / source_size[0]
        scale_y = image.size[1] / source_size[1]

    for key in COVER_FIELD_ORDER:
        placement = scale_box(placements[key], scale_x, scale_y)
        line_box = rendered_lines[key]
        gap = line_box["y_line"] - placement["bottom"]
        if placement["bottom"] >= line_box["y_line"]:
            raise RuntimeError(f"Cover field '{key}' overlaps the underline in the final PDF.")
        if gap < 1.5 or gap > 10:
            raise RuntimeError(
                f"Cover field '{key}' is not aligned above the underline: gap={gap:.2f}px."
            )


def validate_declaration_page_alignment(
    rendered_page: Path,
    placements: dict[str, dict[str, float]],
    label_boxes: dict[str, dict[str, float]],
    source_size: tuple[int, int],
) -> None:
    with Image.open(rendered_page) as image:
        scale_x = image.size[0] / source_size[0]
        scale_y = image.size[1] / source_size[1]

    for key, spec in DECLARATION_FIELD_SPECS.items():
        placement = scale_box(placements[key], scale_x, scale_y)
        label_box = scale_box(label_boxes[str(spec["anchor_label"])], scale_x, scale_y)
        center_delta = abs(
            ((placement["top"] + placement["bottom"]) / 2)
            - ((label_box["top"] + label_box["bottom"]) / 2)
        )
        if center_delta > 2:
            raise RuntimeError(
                f"Declaration field '{key}' is vertically misaligned: delta={center_delta:.2f}px."
            )
        if placement["left"] <= label_box["right"]:
            raise RuntimeError(f"Declaration field '{key}' overlaps its label region.")


def write_toc_entry(paragraph, style_name: str, title: str, page: str) -> None:
    clear_paragraph(paragraph)
    paragraph.style = style_name
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = ABSTRACT_LINE_SPACING
    run = paragraph.add_run(f"{title}\t{page}")
    set_run_text_font(run, "宋体", Pt(12), bold=False)


def normalize_chapter_headings(doc: Document) -> None:
    chapter_pattern = re.compile(r"^\d+\s+")
    for paragraph in doc.paragraphs:
        if paragraph.style.name != "Heading 1":
            continue
        text = paragraph.text.strip()
        if not text:
            paragraph.style = "Normal"
            continue
        if text in {"结 论", "参考文献", "致 谢"}:
            format_unnumbered_heading(paragraph)
            continue
        if chapter_pattern.match(text):
            clear_paragraph(paragraph)
            run = paragraph.add_run(chapter_pattern.sub("", text))
            run.font.name = "Times New Roman"
            set_east_asia_font(run, "黑体")


def rebuild_toc(doc: Document, toc_pages: dict[str, str] | None = None) -> None:
    title_para = None
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == "目  录":
            title_para = paragraph
            break
    if title_para is None:
        raise RuntimeError("Failed to locate TOC title paragraph.")

    toc_end = None
    search_started = False
    for paragraph in doc.paragraphs:
        if paragraph._element is title_para._element:
            search_started = True
            continue
        if search_started and "w:sectPr" in paragraph._element.xml:
            toc_end = paragraph
            break
    if toc_end is None:
        raise RuntimeError("Failed to locate TOC section end.")

    remove_section_break(title_para)
    current = title_para._element.getnext()
    toc_end_element = toc_end._element
    while current is not None and current is not toc_end_element:
        nxt = current.getnext()
        current.getparent().remove(current)
        current = nxt

    format_toc_title(title_para)
    resolved_pages = toc_pages or {title: "" for _, title in TOC_BLUEPRINT}
    for style_name, title in TOC_BLUEPRINT:
        paragraph = toc_end.insert_paragraph_before()
        write_toc_entry(paragraph, style_name, title, resolved_pages.get(title, ""))


def find_table_containing_text(doc: Document, needle: str):
    for table in doc.tables:
        text = "\n".join(cell.text for row in table.rows for cell in row.cells)
        if needle in text:
            return table
    return None


def update_body_content(doc: Document) -> None:
    replacements = {
        "随着金融市场数据规模持续增长":
            "随着金融市场数据规模持续增长，依赖人工经验识别市场热点行业和龙头企业的方法已难以同时满足研究效率、信息时效性和结果可解释性的要求。针对毕业设计任务书提出的热门行业识别与龙头股遴选需求，本文以现有量化研究平台中的行业热度子系统为研究对象，围绕数据获取、行业评分、个股筛选和可视化展示四个环节，设计并实现了一套面向 A 股场景的行业热度分析与龙头股遴选原型系统。",
        "系统采用前后端分离架构。前端基于 React 与 Ant Design 构建行业热力图":
            "系统采用前后端分离架构。前端围绕行业热力图、行业排行榜、行业趋势与龙头股详情构建研究界面；后端基于 FastAPI 提供热门行业、行业成分股、热力图、行业趋势和龙头股等核心接口；分析层使用 Python 完成数据清洗、标准化和综合评分。数据侧采用同花顺优先（Tonghuashun-first，以下简称 THS-first）适配策略，以同花顺行业目录和行业摘要作为主入口，并结合 AKShare、新浪财经和腾讯接口进行字段补齐与回退。",
        "在模型设计方面，本文依据项目中已经实现的 IndustryAnalyzer 与 LeaderStockScorer 两个核心模块":
            "在模型设计方面，本文依据项目中已经实现的 IndustryAnalyzer 行业分析器与 LeaderStockScorer 龙头股评分器，分别构建行业热度评分模型和龙头股综合评分模型。其中，行业热度模型综合考虑价格动量、资金承接、交易活跃度和行业波动率代理四个维度；龙头股评分模型则从市值规模、估值水平、盈利能力、成长性、价格动量和交易活跃度六个维度进行综合评价，并针对列表快速筛选场景设计了轻量评分路径。",
        "结合项目保存的行业热力图历史快照数据":
            "结合项目保存的行业热力图历史快照和系统运行结果进行验证，系统能够在给定统计窗口下识别出电子化学品、通信设备、半导体、消费电子和能源金属等阶段性强势行业，并在行业内部进一步筛选出具有代表性的龙头候选股票。结果表明，该系统具备较好的工程可运行性、结果可解释性和可视化展示能力，能够为热门行业持续跟踪与龙头股遴选提供一定参考。",
        "关键词：金融大数据；行业热度识别；龙头股遴选；多源数据；量化研究平台":
            "关键词：金融大数据；热门行业识别；龙头股遴选；多源数据；量化研究平台",
        "With the rapid growth of financial market data":
            "With the rapid growth of financial market data, approaches that rely primarily on manual experience are no longer sufficient for identifying hot industries and representative leader stocks in a timely and interpretable manner. Targeting the graduation project topic of hot industry recognition and leader stock selection, this thesis takes the industry heat subsystem of an existing quantitative research platform as the research object and develops an A-share oriented prototype covering data acquisition, industry scoring, stock screening, and visual presentation.",
        "The system adopts a front end and back end separated architecture. The front end is implemented with React and Ant Design":
            "The system adopts a front end and back end separated architecture. The front end is implemented with React and Ant Design to provide core research views such as industry heatmaps, ranking boards, trend panels, and leader-stock details. The back end is built on FastAPI, while the analytical layer is implemented in Python. At the data layer, a Tonghuashun-first (THS-first) adapter is used to organize industry catalog data, summary snapshots, and money flow information, with AKShare, Sina Finance, and Tencent valuation endpoints serving as complementary and fallback sources.",
        "According to the actual implementation of the project, the thesis summarizes two core models.":
            "According to the actual implementation of the project, the thesis summarizes two core models. The industry heat model evaluates momentum, money flow strength, trading activity, and a volatility proxy, while the leader stock model evaluates market capitalization, valuation, profitability, growth, price momentum, and trading activity. A lightweight fast scoring path is also retained for scenarios where complete financial data are temporarily unavailable.",
        "Historical heatmap snapshots stored in the project show that the system can iden":
            "Historical heatmap snapshots stored in the project show that the system can identify stage-specific strong industries and further screen representative leader-stock candidates within those industries. The results indicate that the proposed prototype maintains reasonable engineering usability, interpretability, and visual readability, and can provide a practical reference for continuous industry tracking.",
        "Keywords: financial big data; industry heat; leader stock selection; quantitative platform":
            "Keywords: financial big data; hot-industry identification; leading stock selection; quantitative research platform",
        "在热门行业识别与龙头股遴选过程中，不同数据字段的量纲差异很大。例如，行业资金流通常以亿元计量，涨跌幅以百分比计量，市值则以十亿或万亿规模计量。如果直接对原始数据进行加权，容易导致尺度较大的变量主导结果。因此，系统在分析阶段首先对指标做数值清洗和标准化处理。":
            "在热门行业识别与龙头股遴选过程中，不同数据字段的量纲差异很大。例如，行业资金流通常以亿元计量，涨跌幅以百分比计量，市值则以十亿或万亿规模计量。如果直接对原始数据进行加权，容易导致尺度较大的变量主导结果。因此，系统在分析阶段首先对指标做数值清洗和标准化处理。[9-10]",
        "行业评分部分主要使用标准化方法将动量、资金流、活跃度和波动率等因子转换到统一尺度后再进行加权；个股评分部分则结合线性归一化、对数标准化和区间裁剪等方式处理市值、成交额、ROE、营收同比和利润同比等指标。对于缺失数据，系统采用中性值或回退值，保证评分流程不因少量字段缺失而中断。":
            "行业评分部分主要使用标准化方法将动量、资金流、活跃度和波动率等因子转换到统一尺度后再进行加权；个股评分部分则结合线性归一化、对数标准化和区间裁剪等方式处理市值、成交额、净资产收益率（Return on Equity, ROE）、营业收入同比和利润同比等指标。对于缺失数据，系统采用中性值或回退值，保证评分流程不因少量字段缺失而中断。[9-10]",
        "根据项目中的实际实现，行业热度模型以价格动量、资金流强度、交易活跃度和行业波动率四个因子为核心，但在工程实现上并非简单依赖抽象的四因子加总。动量因子主要来自 change_pct 或 weighted_change，用于刻画行业近期相对强弱；资金因子主要来自 flow_strength，用于衡量行业吸引资金的能力；活跃度因子优先采用 avg_volume，缺失时退化为 turnover_rate；波动率因子则优先使用行业指数历史收益率计算的真实波动率，不足时再回退到振幅、换手率或涨跌幅代理值。":
            "根据项目中的实际实现，行业热度模型以价格动量、资金流强度、交易活跃度和行业波动率四个因子为核心，但在工程实现上并非简单依赖抽象的四因子加总。动量因子主要来自 change_pct 或 weighted_change，用于刻画行业近期相对强弱；资金因子主要来自 flow_strength，用于衡量行业吸引资金的能力；活跃度因子优先采用 avg_volume，缺失时退化为 turnover_rate；波动率因子则优先使用行业指数历史收益率计算的真实波动率，不足时再回退到振幅、换手率或涨跌幅代理值。这一处理思路与行业轮动研究中强调多维信号综合判断的观点是一致的。[4][7-8]",
        "在实际运行中，并不是所有行业都能稳定获得完整的历史波动率数据。为提高系统鲁棒性，项目在行业分析引擎中设计了多级回退机制：优先使用真实行业指数历史收益率计算波动率；若历史数据不可用，则依次回退到振幅代理、换手率代理和涨跌幅代理。该设计保证了波动率因子在多数场景下都能被估算出来。":
            "在实际运行中，并不是所有行业都能稳定获得完整的历史波动率数据。为提高系统鲁棒性，项目在行业分析引擎中设计了多级回退机制：优先使用真实行业指数历史收益率计算波动率；若历史数据不可用，则依次回退到振幅代理、换手率代理和涨跌幅代理。该设计保证了波动率因子在多数场景下都能被估算出来，也与金融数据分析中强调缺失容错和稳健处理的实现原则相一致。[9-10]",
        "此外，系统还使用 K-Means 聚类对行业进行辅助划分，并通过轮廓系数自动选择较优聚类数。需要说明的是，聚类分析在本系统中主要承担辅助解释角色，用于观察若干行业是否共同形成热点簇，而不是替代行业热度综合评分本身。":
            "此外，系统还使用 K-Means 聚类方法对行业进行辅助划分，并通过轮廓系数自动选择较优聚类数。需要说明的是，聚类分析在本系统中主要承担辅助解释角色，用于观察若干行业是否共同形成热点簇，而不是替代行业热度综合评分本身。[10]",
        "龙头股评分模型的目标是在已识别出的热门行业或指定行业内部，从多只成分股中筛选更具代表性的股票。结合项目中的实际实现，系统从六个维度构建综合评分模型：市值规模、估值水平、盈利能力、成长性、价格动量和交易活跃度。":
            "龙头股评分模型的目标是在已识别出的热门行业或指定行业内部，从多只成分股中筛选更具代表性的股票。结合项目中的实际实现，系统从六个维度构建综合评分模型：市值规模、估值水平、盈利能力、成长性、价格动量和交易活跃度。相关评价维度与既有龙头企业识别和财务可视化研究中常见的分析口径基本一致。[3][5-6]",
        "其中，市值规模通过对数标准化反映企业体量；估值水平以 PE 处于合理区间得分更高为原则；盈利能力通过 ROE 体现；成长性通过营收同比和利润同比共同体现；价格动量反映短期市场强弱；交易活跃度通过成交额或换手率衡量。设六个维度得分分别为 s1 至 s6，则龙头股综合得分可以表示为：":
            "其中，市值规模通过对数标准化反映企业体量；估值水平以市盈率（Price-to-Earnings Ratio, PE）处于合理区间得分更高为原则；盈利能力通过净资产收益率（Return on Equity, ROE）体现；成长性通过营业收入同比和利润同比共同体现；价格动量反映短期市场强弱；交易活跃度通过成交额或换手率衡量。设六个维度得分分别为 s1 至 s6，则龙头股综合得分可以表示为：[3][5-6]",
        "后端行业接口基于 FastAPI 构建，已经形成较完整的功能集合，包括热门行业列表、行业成分股、热力图、热力图历史、行业趋势、龙头股列表、龙头股详情以及偏好配置等核心接口。聚类与轮动接口在系统中主要承担辅助分析职责，偏好接口还支持按 profile 读写与导入导出，便于前端在不同研究配置之间切换。":
            "后端行业接口基于 FastAPI 构建，已经形成较完整的功能集合，包括热门行业列表、行业成分股、热力图、热力图历史、行业趋势、龙头股列表、龙头股详情以及偏好配置等核心接口。聚类与轮动接口在系统中主要承担辅助分析职责，偏好接口还支持按 profile 读写与导入导出，便于前端在不同研究配置之间切换。该接口组织方式与前文围绕多源金融数据处理和 Python 分析工具链构建研究系统的技术基础相衔接。[1-2][9-10]",
        "技术路线方面，系统采用前后端分离架构。前端使用 React 实现可视化交互":
            "技术路线方面，系统采用前后端分离架构。前端使用 React 实现热力图、排行榜和详情弹窗等核心可视化交互，后端使用 FastAPI 暴露行业分析接口，分析层使用 pandas、NumPy 与 scikit-learn 等库完成数据清洗、标准化、综合评分与辅助聚类分析，数据层以 THS-first 适配器作为行业主入口，并结合 AKShare、新浪和腾讯接口完成补数与回退。研究路线遵循“数据获取、预处理、指标计算、综合评分、结果展示、结果验证”的闭环思路。",
        "多源数据采集是本系统实现的基础。结合项目代码，系统通过数据提供器工厂统一管理不同来源的数据接口。":
            "多源数据采集是本系统实现的基础。结合项目代码，行业子系统以 THS-first 适配策略组织不同来源的数据接口。对于 A 股行业分析任务，同花顺相关接口优先提供行业目录、行业摘要、资金流与领涨股等主数据，AKShare 负责补齐行业元数据、成分股、估值与财务信息，新浪财经和腾讯接口则承担部分回退与字段补充职责。",
        "根据项目中的实际实现，行业热度模型以价格动量、资金流强度、交易活跃度和行业波动率四个因子为核心。":
            "根据项目中的实际实现，行业热度模型以价格动量、资金流强度、交易活跃度和行业波动率四个因子为核心，但在工程实现上并非简单依赖抽象的四因子加总。动量因子主要来自 change_pct 或 weighted_change，用于刻画行业近期相对强弱；资金因子主要来自 flow_strength，用于衡量行业吸引资金的能力；活跃度因子优先采用 avg_volume，缺失时退化为 turnover_rate；波动率因子则优先使用行业指数历史收益率计算的真实波动率，不足时再回退到振幅、换手率或涨跌幅代理值。这一处理思路与行业轮动研究中强调多维信号综合判断的观点是一致的。[4][7-8]",
        "设行业动量、资金流、活跃度和波动率标准化后的结果分别为 Zm、Zf、Zv 和 Zr":
            "设行业动量、资金流、活跃度和波动率标准化后的结果分别为 Zm、Zf、Zv 和 Zr，则行业横截面原始评分可表示为：",
        "系统采用前后端分离架构，前端使用 React 与 Ant Design 构建行业热力图、行业排行榜、龙头股面板、行业趋势图和轮动对比图等交互页面，后端基于 FastAPI 提供数据接口与分析服务，核心分析逻辑由 Python 完成。数据层以 AKShare 获取 A 股行业分类、行业资金流、行业指数、个股估值和财务数据，同时引入新浪行业适配层作为回退机制，以提升系统运行中的稳定性和可用性。":
            "系统采用前后端分离架构，前端使用 React 与 Ant Design 构建行业热力图、行业排行榜、行业趋势图和龙头股详情等核心页面，后端基于 FastAPI 提供行业分析接口与偏好配置服务，核心分析逻辑由 Python 完成。行业数据层并非简单依赖 AKShare 单一数据源，而是由 THS-first 适配器主导：同花顺行业目录与行业摘要提供主要的行业热度底座，AKShare 负责行业元数据、成分股、财务与历史行情补齐，新浪财经和腾讯估值接口承担回退与字段补充，从而提升行业模块的连通性与抗故障能力。",
        "结合项目真实实现，本文的研究内容主要包括四个方面：第一，对现有量化研究平台中的行业热度子系统进行模块化梳理，明确其在整个平台中的功能定位；第二，围绕行业热度识别建立以动量、资金流、活跃度和波动率为核心的行业评分机制；第三，围绕行业内核心公司筛选建立以规模、估值、盈利、成长、动量和活跃度为核心的龙头股评分机制；第四，对系统的前后端实现、缓存机制、历史快照保存和运行结果进行分析。":
            "结合项目真实实现，本文的研究内容主要包括四个方面：第一，对量化研究平台中的行业热度子系统进行模块化梳理，明确 THS-first 数据适配器、FastAPI 路由层和前端行业页面之间的关系；第二，围绕 IndustryAnalyzer 行业分析器梳理行业热度横截面评分、热力图生成和趋势统计的实现机制；第三，围绕 LeaderStockScorer 龙头股评分器分析完整评分与快照快速评分两条龙头股筛选链路；第四，对系统中的缓存、历史快照和偏好持久化进行工程化总结。",
        "本系统的关键技术包括以下几个方面。第一，后端采用 FastAPI 构建 REST 接口，具备较好的开发效率和接口组织能力。第二，前端采用 React 与 Ant Design 构建仪表盘式交互页面，便于完成热力图、表格、趋势图和弹窗详情的联动。第三，数据分析部分主要使用 pandas、NumPy 和 scikit-learn，实现数据处理、标准化与聚类。第四，数据持久化部分采用 JSON 与 SQLite 相结合的方式，以较低成本实现历史快照和用户偏好的保存。[9-10]":
            "本系统的关键技术包括以下几个方面。第一，后端采用 FastAPI 构建 REST 接口，行业模块通过参数化路由输出热门行业、热力图、成分股、龙头股、趋势和偏好配置等核心数据。第二，前端采用 React 与 Ant Design 实现仪表盘式页面，并结合 Recharts 等组件完成热力图、榜单、趋势和详情弹窗联动。第三，分析部分使用 pandas、NumPy 和 scikit-learn 完成字段清洗、横截面标准化、综合评分和辅助聚类分析。第四，行业子系统的持久化主要依赖 JSON 文件与浏览器 localStorage；其他研究模块即便采用结构化存储方式，也不属于本文行业分析主链路的核心内容。[9-10]",
        "系统总体上采用四层结构：表现层、服务层、分析层和数据层。表现层由 IndustryDashboard 及相关前端组件构成，负责用户交互和图表展示；服务层由 FastAPI 行业接口组成，负责请求接收、参数校验、缓存与响应封装；分析层由 IndustryAnalyzer 和 LeaderStockScorer 组成，负责评分计算和结果整理；数据层则通过数据提供器工厂统一管理 AKShare、新浪等数据源。":
            "系统总体上采用四层结构：表现层、服务层、分析层和数据层。表现层以前端行业主页面、排行榜组件以及行业详情和龙头股详情弹窗为核心，回放和偏好配置组件作为辅助交互存在；服务层由 FastAPI 行业接口组成，负责请求接收、参数校验、缓存与响应封装；分析层由 IndustryAnalyzer 行业分析器和 LeaderStockScorer 龙头股评分器组成，负责评分计算和结果整理；数据层则以 THS-first 适配器为核心，统一协调同花顺、AKShare、新浪和腾讯等数据源。",
        "系统数据流程可概括为：前端页面发起请求后，接口层根据请求类型调用行业分析器或龙头股评分器；分析器从主数据源获取行业与个股原始数据，必要时启用回退数据源；随后对数据进行清洗、字段统一和异常兜底；在标准化后计算行业或个股综合得分；最后将结果封装为统一响应结构返回前端，并在本地缓存和快照文件中记录必要信息。":
            "系统数据流程可概括为：前端页面通过 REST 请求调用行业接口；路由层优先检查端点缓存，随后根据请求类型调用 IndustryAnalyzer 或 LeaderStockScorer；分析层优先从 THS-first 适配器获取行业摘要、行业资金流和成分股快照，必要时再回退到 AKShare、新浪或腾讯接口，并在必要时返回过期缓存以维持结果连续性；在字段统一、异常处理和标准化之后，系统计算行业或个股综合得分，将结果封装为统一响应结构返回前端，并把必要的历史快照与偏好配置持久化到本地文件或浏览器缓存。",
        "项目并未采用复杂的大型数据库方案，而是根据实际需求选择轻量化存储方式。行业偏好配置和热力图历史快照保存为 JSON 文件，便于快速读写和迁移；研究工作台任务则使用 SQLite 持久化，以支持更稳定的结构化存储。":
            "项目并未为行业子系统引入独立数据库，而是根据真实实现选择轻量化存储方式。后端将热力图历史快照写入 data/industry/heatmap_history.json，将行业偏好按 profile 写入 data/industry_preferences/<profile>.json，并把龙头股财务缓存保存在 cache/financial_cache.json；前端还将热力图回放快照和当前选中状态保存到浏览器 localStorage 中。其他研究模块若采用结构化存储方式，也不影响本文行业子系统以文件缓存和浏览器状态为主的实现路径。",
        "系统的数据层以数据提供器工厂为统一入口，对不同数据源进行注册、管理与切换。对于行业分析任务，AKShare 是主数据源，承担行业分类、行业资金流、行业指数、成分股、估值与财务信息的获取工作；新浪行业适配层则负责在部分接口失败或字段不完整时进行补充。":
            "系统的数据层以 THS-first 适配器作为行业分析主入口。该适配器优先利用同花顺行业目录与行业一览表获取行业名称、涨跌幅、资金流与领涨股等主数据，再由 AKShare 补齐行业元数据、成分股、财务与历史 K 线，新浪财经接口承担行业列表、成分股与行情补充，腾讯财经接口则补单股估值核心字段。",
        "这种设计使系统形成了主数据源优先、备用数据源兜底的运行模式。对于面向真实市场数据的应用场景而言，这种容错机制比单纯依赖单一接口更具工程可用性。":
            "这种设计使系统形成了“同花顺主导、AKShare 增强、新浪与腾讯补充”的运行模式。对于面向真实市场数据的行业研究场景而言，多源适配比依赖单一接口更具工程可用性，也更能覆盖行业目录、资金流、成分股、估值和历史走势等不同类型的数据需求。",
        "行业分析模块的核心类为 IndustryAnalyzer。该模块负责完成行业资金流分析、行业动量计算、行业波动率估计、热门行业排序、热力图数据生成、行业趋势分析和聚类分析等任务。为提高效率，模块内部设置了 30 分钟缓存，避免在短时间内重复计算相同结果。":
            "行业分析模块的核心类为 IndustryAnalyzer。该模块负责行业资金流分析、行业动量计算、波动率估计、热门行业排序、热力图数据生成和行业趋势统计等任务，聚类与轮动分析则作为辅助能力存在。其实现中优先采用快速路径：直接复用行业资金流与横截面摘要结果，避免逐行业拉取全部成分股；在此基础上，再对 change_pct 或 weighted_change、flow_strength、avg_volume 或 turnover_rate 以及 industry_volatility 做标准化，并按默认权重计算原始分数。得到原始分数后，系统还会进一步压缩到约 20 至 95 的展示区间，以减少样本集中时 0 分和 100 分贴边的问题。为提高效率，分析器内部设置了 30 分钟缓存，避免在短时间内重复计算相同结果。",
        "在接口调用链路中，热门行业接口首先获取行业分类列表，再根据行业资金流和动量数据生成横截面数据表，随后计算综合得分并输出行业排行榜。热力图接口则在此基础上增加尺寸映射和颜色映射所需字段。趋势接口会进一步拉取行业指数序列并汇总行业内涨跌分布。":
            "在接口调用链路中，热门行业接口首先获取行业分类列表，再根据行业资金流、涨跌幅、换手率和波动代理生成横截面数据表，随后计算综合得分并输出行业排行榜。热力图接口则在此基础上增加尺寸映射、颜色映射和市值来源标签所需字段；当真实市值缺失时，系统还会退化到成交额、资金流或成分股数量代理。趋势接口会进一步拉取行业指数序列、汇总成分股涨跌分布，并返回覆盖率与降级说明字段。",
        "龙头股评分模块的核心类为 LeaderStockScorer。该模块负责对单只股票评分、在行业内部进行排名、生成龙头股列表以及输出评分拆解与原始数据。模块同时支持基于完整财务数据的完整评分和基于快照数据的轻量快速评分。":
            "龙头股评分模块的核心类为 LeaderStockScorer，即系统中的龙头股评分器。该模块负责对单只股票评分、在行业内部进行排名、生成龙头股列表以及输出评分拆解与原始数据。模块同时支持两条链路：其一是基于估值与财务数据的完整评分，其二是直接复用行业成分股快照的轻量快速评分。前者用于龙头股详情和深度研究，后者用于行业列表页中的快速筛选。",
        "在实现层面，模块通过缓存机制对财务数据进行暂存，避免在批量评分过程中重复请求；同时使用统一的原始数据结构将估值、财务和行情字段映射为评分输入，使不同评分路径共享同一套维度计算逻辑。":
            "在实现层面，模块会将财务数据以 24 小时缓存方式保存到 financial_cache.json，避免在批量评分过程中重复请求；同时使用统一的 raw_data 原始字段结构映射市值、市盈率（Price-to-Earnings Ratio, PE）、净资产收益率（Return on Equity, ROE）、营收同比、利润同比、涨跌幅和成交额等数据，使完整评分与快速评分共享同一套六维评分逻辑。对于快速评分暂时无法获得的净资产收益率与成长指标，系统采用中性分处理，以兼顾结果稳定性与页面响应效率。",
        "在实现层面，模块会将财务数据以 24 小时 TTL 缓存到 financial_cache.json，避免在批量评分过程中重复请求；同时使用统一的 raw_data 原始字段结构映射市值、PE、ROE、营收同比、利润同比、涨跌幅和成交额等数据，使完整评分与快速评分共享同一套六维评分逻辑。对于快速评分暂时无法获得的 ROE 与增长指标，系统采用中性分处理，以换取页面响应效率。":
            "在实现层面，模块会将财务数据以 24 小时缓存方式保存到 financial_cache.json，避免在批量评分过程中重复请求；同时使用统一的 raw_data 原始字段结构映射市值、市盈率（Price-to-Earnings Ratio, PE）、净资产收益率（Return on Equity, ROE）、营收同比、利润同比、涨跌幅和成交额等数据，使完整评分与快速评分共享同一套六维评分逻辑。对于快速评分暂时无法获得的净资产收益率与成长指标，系统采用中性分处理，以兼顾结果稳定性与页面响应效率。",
        "后端行业接口基于 FastAPI 构建，已经形成较完整的功能集合，包括热门行业列表接口、行业成分股接口、热力图接口、热力图历史接口、行业趋势接口、行业聚类接口、行业轮动接口、龙头股列表接口、龙头股详情接口以及偏好配置接口。":
            "后端行业接口基于 FastAPI 构建，已经形成较完整的功能集合，包括热门行业列表、行业成分股、热力图、热力图历史、行业趋势、龙头股列表、龙头股详情以及偏好配置等核心接口。聚类与轮动接口在系统中主要承担辅助分析职责，偏好接口还支持按 profile 读写与导入导出，便于前端在不同研究配置之间切换。该接口组织方式与前文围绕多源金融数据处理和 Python 分析工具链构建研究系统的技术基础相衔接。[1-2][9-10]",
        "为了提升接口稳定性，接口层额外设置了结果缓存，并在必要时使用过期缓存作为兜底返回。这一设计能够在外部数据源短暂波动时尽量保持页面可用。":
            "为了提升接口稳定性，路由层额外维护了 3 分钟端点缓存、30 分钟 parity 缓存以及完整版成分股异步构建机制，并在必要时返回过期缓存作为补充保障。热力图历史数据还设置了最多 48 条记录和 2MB 文件大小上限，从而在外部数据源短暂波动时尽量保持页面可用。",
        "行业热度页面由 IndustryDashboard 及其子组件实现。页面集成热力图、行业排行榜、行业趋势面板、龙头股面板、轮动分析图、预警面板、观察列表和评分雷达图弹窗等功能。用户可以通过点击行业名称进入详细面板，进一步查看该行业的趋势、成分股和龙头股。":
            "行业热度页面由 IndustryDashboard 及其子组件实现。页面以热力图、行业排行榜、行业趋势面板和龙头股面板为核心，同时保留回放与偏好配置等辅助功能。用户可以通过点击行业名称进入详细面板，进一步查看该行业的趋势、成分股和龙头股。",
        "与传统静态报表不同，本系统更强调交互式研究体验。热力图支持不同颜色维度和尺寸维度切换，排行榜支持多字段排序，详情页支持评分拆解和价格走势查看。这种设计更加符合行业分析工作的实际使用场景。":
            "与传统静态报表不同，本系统更强调交互式研究体验。热力图支持不同颜色维度、尺寸维度和市值来源筛选，排行榜支持回看窗口、波动过滤和排序口径切换，详情页支持评分拆解和价格走势查看；历史快照回放与偏好配置则作为研究辅助能力存在。这种设计更加符合行业分析工作的实际使用场景。",
        "项目中已经实现行业偏好持久化功能，支持保存观察行业、已保存视图和告警阈值。默认告警条件覆盖行业得分、涨跌幅、资金净流入和波动率等指标。该功能说明系统并非一次性分析工具，而是支持持续跟踪的研究原型。":
            "项目中已经实现行业偏好持久化功能，用于保存观察行业和研究配置等辅助信息。后端以 profile 维度将观察行业、研究视图和阈值配置写入 JSON 文件，前端则把热力图回放快照和当前选中状态保存在浏览器本地存储中。上述能力主要服务于持续跟踪与复盘分析，而不是改变行业热度识别与龙头股遴选这一核心研究流程。",
        "此外，系统还会记录热力图历史快照，为后续对比分析提供基础。研究工作台模块则为进一步沉淀分析结果和任务状态提供了可能。":
            "此外，系统后端还会持续记录热力图历史快照，为后续对比分析和前端热力图回放提供基础；这些快照不仅保存涨跌幅，还保留 total_score、资金流、换手率、市值来源和估值字段，便于后续复盘不同时间点的行业状态。",
        "本文的系统测试基于项目当前本地环境展开。系统采用 Python 作为后端分析语言，FastAPI 作为接口框架，React 作为前端框架，数据源主要为 AKShare，并结合项目中保存的历史快照进行结果验证。由于本课题属于毕业设计性质，测试重点放在功能闭环、数据连通性、结果可解释性和页面展示效果，而非面向生产环境的极限压测。":
            "本文的系统测试基于项目当前本地环境展开。系统采用 Python 作为后端分析语言，FastAPI 作为接口框架，React 作为前端框架，行业主数据经由 THS-first 适配器提供，并在需要时调用 AKShare、新浪与腾讯接口补齐字段，同时结合项目中保存的历史快照进行结果验证。除样例快照分析外，项目还保留了行业评分逻辑、适配器映射、偏好服务与前端行业页面的测试脚本，用于验证评分逻辑、回退逻辑和页面交互闭环。由于本课题属于毕业设计性质，测试重点放在功能闭环、数据连通性、结果可解释性和页面展示效果，而非面向生产环境的极限压测。",
        "为了增强论文结果分析的真实性，本文直接选取项目中已经保存的热力图历史快照作为样本。根据 2026 年 4 月 11 日 22:59:20 保存的五日窗口行业热力图数据，系统输出的前十个行业结果如表 6.2 所示。":
            "为了增强论文结果分析的真实性，本文直接选取项目中已经保存的热力图历史快照作为样本。选择 2026 年 4 月 11 日 22:59:20 保存的五日窗口快照，是因为该样本由系统自动记录、字段较为完整，并同时覆盖综合得分、涨跌幅、资金流和换手率等核心指标，便于展示系统在固定统计窗口下的横截面输出。表 6.2 展示了该样本时点按热力图默认顺序截取的前十个行业横截面字段，包括 total_score、5 日涨跌幅、资金流和换手率等信息。需要说明的是，这里给出的结果用于样例分析，并不构成完整回测结论。",
        "从表 6.2 可以看出，通信设备、半导体、电子化学品、消费电子、光学光电子和自动化设备等科技成长风格行业在该时间窗口内表现较强，说明系统能够较好识别当期市场热点方向。值得注意的是，综合得分并不完全等同于涨跌幅。例如，能源金属行业在五日涨跌幅较高的同时，还表现出正向资金流，因此综合得分达到较高水平；而部分行业虽然涨幅较大，但资金流为负，说明系统在排序时并非只依赖价格因子。":
            "从表 6.2 可以看出，电子化学品、通信设备、其他电子、元件、半导体和消费电子等电子科技链行业在该时间窗口内整体表现较强，说明样本时点的市场热点明显集中在成长风格板块。值得注意的是，表中多数科技行业虽然 5 日涨幅较高，但资金流仍为负；只有能源金属表现出较明显的正向资金流。这说明系统生成的 total_score 并不简单等同于单一资金因子或单一涨跌幅，而是横截面标准化后多个因子共同作用的结果。",
        "这一结果与系统的模型设计是一致的。行业热度评分同时考虑动量、资金流、活跃度和波动率，因此能够较好平衡短期趋势与资金确认。对于毕业设计来说，这种结果不仅能够支持论文中的实验分析，也便于在答辩中解释评分逻辑。":
            "这一结果与系统代码中的模型设计是一致的。热力图快照不仅保留 value 和 total_score，还同步记录资金流、换手率、行业波动率和市值来源等字段，因此能够较好平衡短期趋势、资金确认和数据来源差异。需要强调的是，total_score 反映的是样本时点的多因子横截面综合状态，并不直接构成对未来收益的预测。对于毕业设计来说，这种结果既支持论文中的样例分析，也便于在答辩中解释为什么某些行业会在资金流暂时偏弱时仍保持较高的综合得分。",
        "同时，系统仍存在一些不足。首先，行业评分和龙头股评分的权重主要依据工程经验设定，尚未通过更系统的学习或优化方法进行自适应调整。其次，系统当前对政策文本、新闻舆情和产业链关系等数据的融合仍不充分。再次，系统的“实时性”主要体现为准实时刷新与历史快照持续积累，并非高频交易系统意义上的毫秒级实时。最后，行业识别结果与未来收益之间的定量回测仍可进一步强化。":
            "同时，系统仍存在一些不足。首先，行业评分和龙头股评分的权重主要依据工程经验设定，尚未通过更系统的学习或优化方法进行自适应调整。其次，虽然 THS-first 适配器提高了整体可用性，但部分行业成分股在实际运行中仍可能触发降级路径，只能返回不完整明细或代理字段。再次，行业模块当前主要依赖 REST 拉取、缓存和历史快照实现准实时更新，而非基于持续流式订阅的实时分析系统。最后，行业识别结果与后续收益表现之间的定量回测，以及前瞻性验证工作，仍有进一步强化空间。",
        "本文以现有量化研究平台中的行业热度子系统为依托，围绕热门行业识别与龙头股遴选这一毕业设计主题，完成了系统分析、模型设计、工程实现和结果总结。通过梳理项目中的数据提供器、行业分析引擎、龙头股评分器、FastAPI 接口层和 React 仪表盘页面，本文构建了一条从多源数据获取到可视化展示的完整研究主线。":
            "本文以现有量化研究平台中的行业热度子系统为依托，围绕热门行业识别与龙头股遴选这一毕业设计主题，完成了系统分析、模型设计、工程实现和结果总结。通过梳理项目中的 THS-first 数据适配器、行业分析器、龙头股评分器、后端接口层以及前端行业仪表盘页面，本文构建了一条从多源数据获取到可视化展示的完整研究主线。",
        "在模型方面，本文总结并实现了以动量、资金流、活跃度和波动率为核心的行业热度评分模型，以及以规模、估值、盈利、成长、动量和活跃度为核心的龙头股综合评分模型；在实现方面，本文说明了多源数据回退、缓存机制、历史快照、偏好持久化和前端可视化协同工作方式；在结果方面，本文结合项目保存的历史热力图快照展示了系统对阶段性强势行业的识别能力。":
            "在模型方面，本文总结并实现了以动量、资金流、活跃度和波动率代理为核心的行业热度评分模型，以及以规模、估值、盈利、成长、动量和活跃度为核心的龙头股综合评分模型；在实现方面，本文说明了 THS-first 多源数据回退、分析层与路由层缓存、热力图历史快照和按 profile 偏好持久化的协同工作方式；在结果方面，本文结合项目保存的历史热力图快照展示了系统对阶段性强势行业的识别能力。",
        "未来工作可以从以下方向展开：第一，引入舆情、政策文本和产业链数据，增强行业识别的前瞻性；第二，采用熵权法、层次分析法或机器学习方法优化权重；第三，完善行业识别结果的回测验证；第四，强化与研究工作台的联动，形成更完整的研究闭环。总体而言，本文所完成的工作已经能够满足本科毕业设计对理论分析、系统实现与应用展示的综合要求。":
            "未来工作可以从以下方向展开：第一，引入舆情、政策文本和产业链数据，增强行业识别的前瞻性；第二，采用熵权法、层次分析法或机器学习方法优化权重；第三，完善行业识别结果与后续收益表现之间的回测验证；第四，继续提高成分股映射完整度并优化核心行业分析链路。总体而言，本文所完成的工作已经能够满足本科毕业设计对理论分析、系统实现与应用展示的综合要求。",
        "该公式中的权重与项目代码中的默认权重保持一致":
            "该公式中的权重与项目代码中的默认权重保持一致。需要强调的是，公式（4.2.1）对应的是横截面原始评分，并不直接等同于前端最终展示值。系统在完成加权之后，会进一步调用分数压缩函数，把结果统一映射到约 20 至 95 的展示区间，以减少样本集中时 0 分和 100 分贴边的问题。因此，排行榜和热力图中的 total_score 更适合用于相对比较，而不应被理解为行业的绝对评价值。",
        "此外，系统还使用 K-Means 聚类对行业进行辅助划分":
            "此外，系统还使用 K-Means 聚类对行业进行辅助划分，并通过轮廓系数自动选择较优聚类数。需要说明的是，聚类分析在本系统中主要承担辅助解释角色，用于观察若干行业是否共同形成热点簇，而不是替代行业热度综合评分本身。",
        "该模型突出盈利和成长的权重":
            "从工程实现看，完整评分链路会在统一的 raw_data 结构上聚合估值、财务和行情字段，并按六维权重体系生成总分。由于盈利能力和成长性权重相对较高，该模型更适合作为行业内核心标的筛选依据；与单纯按涨幅排序相比，它能够减少短期情绪对结果的扰动。",
        "快速评分模式适合用于行业热度页中的候选股预筛选":
            "快速评分模式适合用于行业热度页中的候选股预筛选，而完整评分模式适合用于龙头股详情与深度研究。两条链路共享同一套权重体系和输出结构，只是在输入字段完备性上存在差异，因此系统能够在保证评分口径一致的前提下兼顾响应效率与分析完整性。",
        "从工程角度看，本系统具有以下几个明显特征":
            "从工程角度看，本系统具有以下几个明显特征。第一，系统以 THS-first 适配器为主入口，并通过 AKShare、新浪和腾讯形成多源补齐链路，保证了行业数据的连通性。第二，行业评分与龙头股评分同时引入快速评分链路与多级缓存机制，在保证响应速度的同时尽量维持评分口径一致。第三，系统输出不仅提供综合得分，还提供资金、波动率、市值来源和评分拆解等辅助信息，因此具有较好的结果解释性。",
        "在本次毕业设计与论文写作过程中":
            "在本次毕业设计与论文写作过程中，感谢指导教师沈文辉老师在选题论证、系统设计、论文结构与写作规范等方面给予的耐心指导与帮助。感谢项目开发和调试过程中提供建议与支持的同学和朋友，他们的交流与反馈为系统完善和论文修改提供了重要参考。",
        "同时，也感谢开源社区提供的 Python、FastAPI、React、AKShare 等工具和框架":
            "同时，也感谢开源社区提供的 Python、FastAPI、React、AKShare、scikit-learn 等工具与框架，使本课题能够在真实工程环境中完成设计、实现与验证。最后，感谢学校和学院为毕业设计工作提供的学习环境与资源支持。"
    }

    for old_text, new_text in replacements.items():
        replace_first_paragraph_starting_with(doc, old_text, new_text)

    replace_section_body(
        doc,
        "1.2 国内外研究现状",
        "1.3 研究内容与技术路线",
        [
            "1.2.1 金融大数据分析研究现状",
            "在金融大数据分析方面，已有研究已经形成较为成熟的方法基础。孟小峰等对大数据管理技术进行了系统梳理，指出容量、速度、多样性和真实性构成了金融数据处理的核心挑战。[1] 张晓琳等从风险监测与预警角度讨论了多源金融数据融合的必要性。[2] 在工具层面，McKinney 对 Pandas、NumPy 等 Python 数据分析工具的总结，以及 scikit-learn 在标准化、聚类和模式识别方面提供的统一接口，为金融数据清洗、转换和建模提供了成熟的技术支撑。[9-10]",
            "1.2.2 行业轮动与热点识别研究现状",
            "在行业轮动与热点识别方面，国内外研究普遍强调多因子综合判断的重要性。相关工作通常将宏观周期、风格切换、价格动量、资金流向和交易活跃度等变量结合起来，用于识别阶段性强势行业。[4][7-8] 这说明热门行业识别并不能依赖单一指标，而应在横截面上综合考察价格、资金与活跃度等多维信号。",
            "1.2.3 龙头股筛选与系统实现研究现状",
            "在龙头股识别方面，已有研究更强调基本面与市场表现的结合。企业规模、盈利能力、成长性和估值水平是较常见的评价维度，而近期动量、成交额和资金关注度则用于反映短期市场确认。[3][5-6] 这类研究为本文构建龙头股综合评分模型提供了理论依据，但不少工作更偏向静态评价或策略回测，缺少与真实软件系统的数据链路、缓存机制和可视化展示的紧密结合。",
            "总体来看，现有研究为热门行业识别与龙头股筛选提供了较为丰富的理论与方法基础，但仍存在两点不足：一是较少围绕真实研究系统讨论多源适配、回退链路和结果展示；二是行业识别与龙头股遴选往往被分开讨论，缺少统一的工程闭环。本文的工作重点并不在于构建复杂黑箱模型，而是在真实项目基础上，把多源数据获取、行业评分、龙头股筛选和前端展示整合为一个可运行、可解释的行业研究原型系统。",
        ],
    )

    replace_section_body(
        doc,
        "1.3 研究内容与技术路线",
        "1.4 论文结构安排",
        [
            "结合项目真实实现，本文的研究内容主要包括四个方面：第一，对量化研究平台中的行业热度子系统进行模块化梳理，明确同花顺优先（Tonghuashun-first，THS-first）数据适配器、FastAPI 路由层和前端行业页面之间的关系；第二，围绕行业分析器总结行业热度横截面评分、热力图生成和趋势统计的实现机制；第三，围绕龙头股评分器分析完整评分与快照快速评分两条龙头股筛选链路；第四，对系统中的缓存、历史快照和偏好持久化进行工程化总结，并结合样例快照与测试结果验证其合理性。",
            "技术路线方面，系统采用前后端分离架构。前端使用 React 实现热力图、排行榜、趋势面板与详情弹窗等核心交互，后端使用 FastAPI 暴露热门行业、成分股、热力图、趋势分析、龙头股与偏好配置等接口，分析层使用 pandas、NumPy 与 scikit-learn 完成字段清洗、标准化、综合评分与辅助聚类分析，数据层则以同花顺优先（THS-first）适配器作为行业主入口，并结合 AKShare、新浪和腾讯接口完成补数与回退。整体研究路线遵循“数据获取、字段统一、指标计算、综合评分、结果展示、结果验证”的研究流程。",
        ],
    )

    replace_section_body(
        doc,
        "3.4 系统总体架构设计",
        "3.5 数据流程设计",
        [
            "如图 3.1 所示，系统总体上采用四层结构：表现层、服务层、分析层和数据层。表现层以前端行业主页面、热力图、排行榜组件以及行业详情和龙头股详情弹窗为核心，回放、观察列表和偏好配置等功能作为辅助研究交互存在；服务层由 FastAPI 行业接口组成，负责请求接收、参数校验、缓存与响应封装；分析层由行业分析器和龙头股评分器组成，负责行业综合评分、龙头股筛选和评分拆解；数据层则以 THS-first 适配器为核心，统一协调同花顺、AKShare、新浪和腾讯等数据源。",
            "这种分层方式使系统既能保持接口职责清晰，又能兼顾研究系统的灵活性。表现层只关注图表交互与结果展示，服务层负责对分析结果进行统一组织，分析层专注评分计算与结果解释，数据层则通过名称映射、节点回退、符号缓存和过期缓存保障等机制提升数据可用性。即使未来更换局部数据源或调整评分权重，也主要集中在分析层和数据层完成，不需要整体重写前端页面。",
            "与仅依赖单一接口的数据抓取脚本相比，该架构更适合支撑准实时行业研究场景。一方面，多源适配提高了行业目录、资金流、成分股、估值和历史走势等字段的覆盖率；另一方面，缓存与分层封装保证了页面响应速度和结果口径的一致性，使系统能够在真实研究环境中持续运行并为论文分析提供稳定样本。",
        ],
    )

    replace_section_body(
        doc,
        "6.5 对毕业设计目标的达成情况",
        "结 论",
        [
            "对照任务书要求，如表 6.3 所示，本文在理论分析、模型设计和系统实现三个层面均完成了核心目标。任务书提出的“理解金融大数据挖掘相关技术与研究现状”“掌握行业分析与企业筛选方法”“编程设计软件系统实现重点行业和关键企业跟踪”三项要求，均已在论文和系统中得到对应落实。",
            "虽然系统并未采用最初设想中的大规模分布式基础设施，而是采用更贴合项目实际的轻量化架构，但从工程完整性、研究闭环和论文展示效果来看，这种实现路径更符合本科毕业设计对可运行性、可解释性和可展示性的综合要求。",
        ],
    )

    replace_section_body(
        doc,
        "5.1 数据提供器与回退机制实现",
        "5.2 行业分析模块实现",
        [
            "行业子系统的数据入口封装在 SinaIndustryAdapter 中，其核心思想是同花顺优先（Tonghuashun-first，THS-first）。具体而言，同花顺接口优先提供行业目录、行业摘要、行业资金流、领涨股和行业指数等主数据；AKShare 负责补齐行业元数据、成分股、估值、财务与历史行情；新浪财经接口承担行业列表、成分股和部分实时行情的补充获取；腾讯财经接口则在单股估值字段缺失时提供补充。通过这种职责分工，系统形成了“同花顺主导、AKShare 增强、新浪与腾讯补充”的运行模式，该思路与金融大数据场景下强调多源异构数据整合与交叉验证的研究认识是一致的。[1-2]",
            "为了处理不同数据源之间的口径差异，适配层实现了较完整的名称映射与回退机制。例如，系统维护了新浪行业名称到同花顺行业名称的映射表，并针对部分行业准备了专门的节点映射与代理节点映射；当主路径失败时，系统会尝试使用反向名称映射、符号缓存和过期缓存继续补全结果，从而提高行业目录与成分股匹配的稳定性。",
            "这种数据适配设计并非对多个接口的简单叠加，而是围绕行业研究任务进行有组织的字段整合。对论文研究而言，它保证了行业热度计算所需的行业摘要、资金流、成分股、估值和财务字段可以在统一的数据入口下被消费；对工程实现而言，它显著降低了外部数据源短时波动对系统可用性的影响，为后续行业评分、龙头股筛选和热力图历史快照记录提供了稳定的数据底座。",
        ],
    )

    replace_section_body(
        doc,
        "结 论",
        "参考文献",
        [
            "本文以现有量化研究平台中的行业热度子系统为依托，围绕热门行业识别与龙头股遴选这一毕业设计主题，完成了系统分析、模型设计、工程实现和结果总结。通过梳理项目中的 THS-first 数据适配器、行业分析器、龙头股评分器、后端接口层以及前端行业仪表盘页面，本文构建了一条从多源数据获取到可视化展示的完整研究主线。",
            "在模型方面，本文总结并实现了以动量、资金流、活跃度和波动率代理为核心的行业热度评分模型，以及以规模、估值、盈利、成长、动量和活跃度为核心的龙头股综合评分模型。行业热度部分采用横截面标准化与加权合成的方法，并进一步压缩到便于展示的相对得分区间；龙头股部分则同时支持完整评分与快照快速评分两条链路，在保证口径一致的前提下兼顾分析完整性与页面响应效率。",
            "在工程实现方面，本文说明了 THS-first 多源数据回退、分析层与路由层缓存、热力图历史快照、财务缓存以及按 profile 偏好持久化的协同工作方式。第六章结合项目保存的历史热力图快照和测试结果说明，该系统能够识别样本时点的阶段性强势行业，并进一步给出行业内部具有代表性的龙头股候选结果，具有较好的结果解释性与展示性。",
            "未来工作可以从以下方向展开：第一，引入舆情、政策文本和产业链数据，增强行业识别的前瞻性；第二，采用熵权法、层次分析法或机器学习方法优化权重；第三，完善行业识别结果与后续收益表现之间的回测验证；第四，继续提高成分股映射完整度并优化核心行业分析链路。总体而言，本文所完成的工作已经能够满足本科毕业设计对理论分析、系统实现与应用展示的综合要求。",
        ],
    )

    test_heading = next((p for p in doc.paragraphs if p.text.strip() == "6.2 功能测试"), None)
    if test_heading is not None:
        next_element = test_heading._element.getnext()
        intro_text = (
            "功能测试围绕分析逻辑正确性、接口返回稳定性和前端交互闭环三个目标展开。分析层单元测试用于"
            "验证行业热度计算中的快速评分链路、回退逻辑、历史波动率优先级以及龙头股评分逻辑；接口与适配器测"
            "试用于验证 THS-first 适配器、行业偏好服务和龙头股详情接口的数据返回结构；前端端到端测试"
            "用于验证热力图切换、行业搜索、详情弹窗和排行榜交互是否形成完整闭环。表 6.1 对主要测试内"
            "容与结果进行了归纳，结果表明系统核心功能均达到预期。"
        )
        if next_element is not None and next_element.tag == qn("w:p"):
            from docx.text.paragraph import Paragraph

            next_para = Paragraph(next_element, test_heading._parent)
            if next_para.text.strip().startswith("功能测试主要分为三类"):
                replace_paragraph_text(next_para, intro_text)
            elif next_para.text.strip() != "表 6.1 主要功能测试结果":
                replace_paragraph_text(next_para, intro_text)
            else:
                intro_para = insert_paragraph_after(test_heading, intro_text)
                intro_para.paragraph_format.keep_with_next = True

    function_table = find_table_containing_text(doc, "模块")
    if function_table is not None and len(function_table.rows) >= 6 and len(function_table.columns) >= 3:
        function_table.cell(1, 1).text = "获取 THS 行业目录、行业摘要、成分股与财务等数据"
        function_table.cell(1, 2).text = "THS-first 适配器主导，AKShare/Sina/腾讯补齐"
        function_table.cell(2, 2).text = "横截面评分、波动率代理、聚类与轮动分析"
        function_table.cell(4, 2).text = "React 仪表盘、热力图与详情联动"
        function_table.cell(5, 2).text = "按 profile 的 JSON 持久化与浏览器缓存"

    sample_table = find_table_containing_text(doc, "资金流(元)")
    if sample_table is not None and len(sample_table.rows) >= 1 and len(sample_table.columns) >= 5:
        sample_table.cell(0, 0).text = "序号"
        sample_table.cell(0, 4).text = "资金流(亿元)"
        for row_index in range(1, min(len(sample_table.rows), 11)):
            try:
                value = float(sample_table.cell(row_index, 4).text)
                sample_table.cell(row_index, 4).text = f"{value / 1e8:.2f}"
            except ValueError:
                continue

    test_table = find_table_containing_text(doc, "测试项目")
    if test_table is None:
        test_table = find_table_containing_text(doc, "输入或操作")
    if test_table is not None and len(test_table.rows) >= 6 and len(test_table.columns) >= 4:
        test_table.cell(0, 0).text = "测试项目"
        row_values = [
            ("行业评分计算与回退", "运行行业评分逻辑并验证缺失字段回退", "快速评分链路可输出得分，波动率与字段回退生效", "满足预期"),
            ("THS-first 适配器映射", "验证行业名称映射和多源回退", "THS 主源可用，AKShare/Sina/腾讯补齐链路正常", "满足预期"),
            ("龙头股列表与详情接口", "调用龙头股列表与详情接口", "返回综合得分、维度拆解和原始字段", "满足预期"),
            ("偏好配置持久化", "保存并重置观察行业与阈值配置", "JSON 与浏览器状态保持一致", "满足预期"),
            ("前端行业页面闭环", "切换热力图周期、搜索行业并打开详情", "热力图、排行榜与详情弹窗联动正常", "满足预期"),
        ]
        for row_index, values in enumerate(row_values, start=1):
            for col_index, value in enumerate(values):
                test_table.cell(row_index, col_index).text = value

    insert_architecture_figure(doc)
    insert_task_completion_table(doc)


def find_paragraph_by_text(doc: Document, text: str):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise RuntimeError(f"Paragraph not found: {text}")


def find_last_paragraph_by_text(doc: Document, text: str):
    for paragraph in reversed(doc.paragraphs):
        if paragraph.text.strip() == text:
            return paragraph
    raise RuntimeError(f"Paragraph not found: {text}")


def find_previous_drawing_paragraph(doc: Document, anchor_paragraph):
    paragraphs = doc.paragraphs
    anchor_index = None
    for idx, paragraph in enumerate(paragraphs):
        if paragraph._element is anchor_paragraph._element:
            anchor_index = idx
            break
    if anchor_index is None:
        raise RuntimeError(f"Anchor paragraph not found in document: {anchor_paragraph.text}")
    for idx in range(anchor_index - 1, -1, -1):
        paragraph = paragraphs[idx]
        if "w:drawing" in paragraph._element.xml:
            return paragraph
    raise RuntimeError(f"Drawing paragraph not found before: {anchor_paragraph.text}")


def ensure_figure_source_paragraph(caption):
    next_element = caption._element.getnext()
    if next_element is not None and next_element.tag == qn("w:p"):
        from docx.text.paragraph import Paragraph

        next_para = Paragraph(next_element, caption._parent)
        if next_para.text.strip().startswith("图片来源："):
            replace_paragraph_text(next_para, SOURCE_NOTE)
            return next_para
    return insert_paragraph_after(caption, SOURCE_NOTE)


def relayout_figures(doc: Document) -> None:
    body_section = doc.sections[-1]
    max_body_width = body_section.page_width - body_section.left_margin - body_section.right_margin
    TMP_ASSET_DIR.mkdir(parents=True, exist_ok=True)
    for title, config in FIGURE_IMAGES.items():
        caption = find_paragraph_by_text(doc, title)
        image_para = find_previous_drawing_paragraph(doc, caption)
        previous = image_para._element.getprevious()
        if previous is not None and previous.tag == qn("w:p") and 'w:type="page"' in previous.xml:
            previous.getparent().remove(previous)
        source_path = Path(config["path"])
        crop_box = config["crop"]
        output_path = TMP_ASSET_DIR / f"{config['slug']}.png"
        with Image.open(source_path) as img:
            cropped = img.crop(crop_box)
            bordered = ImageOps.expand(cropped, border=4, fill="#C9D4E0")
            bordered.save(output_path)
        clear_paragraph(image_para)
        image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        image_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        image_para.paragraph_format.line_spacing = 1.0
        image_para.paragraph_format.space_before = Pt(0)
        image_para.paragraph_format.space_after = Pt(0)
        image_para.paragraph_format.first_line_indent = Pt(0)
        remove_page_break_before_flag(image_para)
        width_scale = config.get("width_scale", 0.97)
        target_width = Emu(int(max_body_width * width_scale))
        image_para.add_run().add_picture(str(output_path), width=target_width)
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.paragraph_format.first_line_indent = Pt(0)
        caption.paragraph_format.space_before = Pt(3)
        caption.paragraph_format.space_after = Pt(0)
        caption.paragraph_format.keep_with_next = True
        image_para.paragraph_format.keep_with_next = True
        source_para = ensure_figure_source_paragraph(caption)
        source_para.paragraph_format.space_before = Pt(0)
        source_para.paragraph_format.space_after = Pt(3)
        source_para.paragraph_format.keep_with_next = False


def _set_row_no_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def set_table_borders(table, top: bool, header_bottom: bool, bottom: bool) -> None:
    rows = table.rows
    for row_idx, row in enumerate(rows):
        _set_row_no_split(row)
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_borders = tc_pr.find(qn("w:tcBorders"))
            if tc_borders is None:
                tc_borders = OxmlElement("w:tcBorders")
                tc_pr.append(tc_borders)
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                node = tc_borders.find(qn(f"w:{edge}"))
                if node is None:
                    node = OxmlElement(f"w:{edge}")
                    tc_borders.append(node)
                enabled = False
                if edge == "top" and top and row_idx == 0:
                    enabled = True
                elif edge == "bottom" and header_bottom and row_idx == 0:
                    enabled = True
                elif edge == "bottom" and bottom and row_idx == len(rows) - 1:
                    enabled = True
                if enabled:
                    node.set(qn("w:val"), "single")
                    node.set(qn("w:sz"), "8")
                    node.set(qn("w:space"), "0")
                    node.set(qn("w:color"), "000000")
                else:
                    node.set(qn("w:val"), "nil")


def format_data_tables(doc: Document) -> None:
    for table in doc.tables:
        text = "\n".join(cell.text for row in table.rows for cell in row.cells)
        if "姓    名：" in text or "论文题目：" in text:
            continue
        if "行业热力图 / 行业排行榜" in text and "THS 主数据 / AKShare 增强" in text:
            continue
        if "Sindustry" in text or "Sleader" in text:
            continue
        if not table.rows:
            continue
        set_table_borders(table, top=True, header_bottom=True, bottom=True)
        for row_index, row in enumerate(table.rows):
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 else WD_ALIGN_PARAGRAPH.LEFT
                    paragraph.paragraph_format.first_line_indent = Pt(0)
                    paragraph.paragraph_format.line_spacing = Pt(20)
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    for run in paragraph.runs:
                        run.font.bold = row_index == 0
                        set_run_text_font(run, "宋体", Pt(10.5))


def replace_formula_table(doc: Document, table_keyword: str, equation_text: str, equation_number: str) -> None:
    target_table = find_table_containing_text(doc, table_keyword)
    if target_table is None:
        return

    table_element = target_table._element
    parent = target_table._parent
    equation_para = insert_paragraph_before_element(table_element, parent)
    equation_para.style = "Normal"
    equation_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    equation_para.paragraph_format.first_line_indent = Pt(0)
    equation_para.paragraph_format.space_before = Pt(6)
    equation_para.paragraph_format.space_after = Pt(6)
    equation_para.paragraph_format.line_spacing = BODY_LINE_SPACING
    usable_width = doc.sections[-1].page_width - doc.sections[-1].left_margin - doc.sections[-1].right_margin
    equation_para.paragraph_format.tab_stops.add_tab_stop(Emu(int(usable_width / 2)), alignment=WD_TAB_ALIGNMENT.CENTER)
    equation_para.paragraph_format.tab_stops.add_tab_stop(Emu(int(usable_width)), alignment=WD_TAB_ALIGNMENT.RIGHT)
    run = equation_para.add_run(f"\t{equation_text}\t{equation_number}")
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    set_east_asia_font(run, "Times New Roman")
    table_element.getparent().remove(table_element)


def replace_formula_tables(doc: Document) -> None:
    replace_formula_table(
        doc,
        "Sindustry",
        "Sindustry=0.35×Zm+0.35×Zf+0.15×Zv-0.15×Zr",
        "（4.2.1）",
    )
    replace_formula_table(
        doc,
        "Sleader",
        "Sleader=100×(0.20×s1+0.15×s2+0.25×s3+0.20×s4+0.10×s5+0.10×s6)",
        "（4.4.1）",
    )


def normalize_reference_entries(doc: Document) -> None:
    entries = [
        "[1] 孟小峰, 慈祥. 大数据管理：概念、技术与挑战[J]. 计算机学报, 2013, 36(3): 635-652.",
        "[2] 张晓琳, 王吉. 基于大数据的金融风险监测与预警研究[J]. 金融理论与实践, 2022(4): 58-64.",
        "[3] 李明, 刘强. 基于知识图谱的产业链核心企业识别方法研究[J]. 计算机集成制造系统, 2023, 29(5): 1650-1662.",
        "[4] 陈云, 赵刚. 行业轮动策略在量化投资中的应用研究：基于宏观经济数据的实证分析[J]. 投资研究, 2021, 40(9): 45-58.",
        "[5] 王赫. 基于 Python 的上市公司财务数据可视化分析系统的设计与实现[D]. 北京: 北京交通大学, 2020.",
        "[6] 韦韦. 机器学习在金融大数据挖掘中的应用场景探讨[J]. 现代信息科技, 2023, 7(12): 115-118.",
        [
            "[7] Cavalcante R C, Brasileiro R C, Souza V L, et al. Computational Intelligence and Financial Markets: A Survey and Future Directions[J].",
            "Expert Systems with Applications, 2016, 55: 194-211.",
        ],
        [
            "[8] Fisher I, et al. Analytics for Financial Data: Methods and Applications[J].",
            "Journal of Big Data Analytics in Finance, 2022, 3(1): 22-35.",
        ],
        [
            "[9] McKinney W. Python for Data Analysis: Data Wrangling with Pandas, NumPy, and IPython[M].",
            "3rd ed. Sebastopol: O'Reilly Media, 2022.",
        ],
        [
            "[10] Pedregosa F, Varoquaux G, Gramfort A, et al. Scikit-learn: Machine Learning in Python[J].",
            "Journal of Machine Learning Research, 2011, 12: 2825-2830.",
        ],
    ]
    start_index = next(
        (idx + 1 for idx, paragraph in enumerate(doc.paragraphs) if paragraph.text.strip() == "参考文献"),
        None,
    )
    if start_index is None:
        raise RuntimeError("Reference section anchor not found.")
    for offset, entry in enumerate(entries):
        paragraph = doc.paragraphs[start_index + offset]
        clear_paragraph(paragraph)
        if isinstance(entry, str):
            paragraph.add_run(entry)
            continue
        for index, line in enumerate(entry):
            run = paragraph.add_run(line)
            if index < len(entry) - 1:
                run.add_break(WD_BREAK.LINE)


def apply_minor_layout_overrides(doc: Document) -> None:
    target_heading = next(
        (paragraph for paragraph in doc.paragraphs if paragraph.text.strip() == "5.6 偏好配置与持续跟踪实现"),
        None,
    )
    if target_heading is not None:
        target_heading.paragraph_format.page_break_before = True
        target_heading.paragraph_format.keep_with_next = True
        next_element = target_heading._element.getnext()
        if next_element is not None and next_element.tag == qn("w:p"):
            from docx.text.paragraph import Paragraph

            first_para = Paragraph(next_element, target_heading._parent)
            first_para.paragraph_format.keep_with_next = True
            second_element = next_element.getnext()
            if second_element is not None and second_element.tag == qn("w:p"):
                second_para = Paragraph(second_element, target_heading._parent)
                second_para.paragraph_format.keep_together = True

    completion_heading = next(
        (paragraph for paragraph in doc.paragraphs if paragraph.text.strip() == "6.5 对毕业设计目标的达成情况"),
        None,
    )
    if completion_heading is not None:
        completion_heading.paragraph_format.page_break_before = True
        completion_heading.paragraph_format.keep_with_next = True
        next_element = completion_heading._element.getnext()
        if next_element is not None and next_element.tag == qn("w:p"):
            from docx.text.paragraph import Paragraph

            intro_para = Paragraph(next_element, completion_heading._parent)
            intro_para.paragraph_format.keep_with_next = True


def apply_superscript_citations(paragraph) -> None:
    text = paragraph.text
    if not text or "[" not in text or paragraph.text.strip().startswith("["):
        return
    matches = list(re.finditer(r"\[[0-9,\-，]+\]", text))
    if not matches:
        return
    clear_paragraph(paragraph)
    cursor = 0
    for match in matches:
        if match.start() > cursor:
            normal_run = paragraph.add_run(text[cursor:match.start()])
            format_body_run(normal_run)
        citation_run = paragraph.add_run(match.group(0))
        format_body_run(citation_run)
        citation_run.font.superscript = True
        cursor = match.end()
    if cursor < len(text):
        tail_run = paragraph.add_run(text[cursor:])
        format_body_run(tail_run)


def format_paragraphs(doc: Document) -> None:
    abstract_index = next(
        (idx for idx, paragraph in enumerate(doc.paragraphs) if paragraph.text.strip() == "摘  要"),
        0,
    )
    english_abstract_index = next(
        (idx for idx, paragraph in enumerate(doc.paragraphs) if paragraph.text.strip() == "ABSTRACT"),
        abstract_index,
    )
    toc_index = next(
        (idx for idx, paragraph in enumerate(doc.paragraphs) if paragraph.text.strip() == "目  录"),
        english_abstract_index,
    )
    for idx, paragraph in enumerate(doc.paragraphs):
        if idx < abstract_index:
            continue
        text = paragraph.text.strip()
        style_name = paragraph.style.name if paragraph.style is not None else ""

        if not text:
            continue
        if text == "摘  要":
            clear_paragraph(paragraph)
            run = paragraph.add_run("摘  要")
            paragraph.style = "Normal"
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = ABSTRACT_LINE_SPACING
            set_run_text_font(run, "黑体", Pt(18), bold=True)
            continue
        if text == "ABSTRACT":
            clear_paragraph(paragraph)
            run = paragraph.add_run("ABSTRACT")
            paragraph.style = "Normal"
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = ABSTRACT_LINE_SPACING
            set_run_text_font(run, "Times New Roman", Pt(18), bold=True)
            continue
        if text == "目  录":
            format_toc_title(paragraph)
            continue
        if style_name in {"toc 1", "toc 2"}:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = ABSTRACT_LINE_SPACING
            for run in paragraph.runs:
                set_run_text_font(run, "宋体", Pt(12), bold=False)
            continue
        if style_name == "Heading 1":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(12)
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = BODY_LINE_SPACING
            for run in paragraph.runs:
                set_run_text_font(run, "黑体", Pt(18), bold=True)
            continue
        if style_name == "Heading 2":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = BODY_LINE_SPACING
            for run in paragraph.runs:
                set_run_text_font(run, "黑体", Pt(14), bold=True)
            continue
        if style_name == "Heading 3":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(3)
            paragraph.paragraph_format.line_spacing = BODY_LINE_SPACING
            for run in paragraph.runs:
                set_run_text_font(run, "黑体", Pt(12), bold=True)
            continue
        if text in {"结 论", "参考文献", "致 谢"}:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(12)
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = BODY_LINE_SPACING
            for run in paragraph.runs:
                set_run_text_font(run, "黑体", Pt(18), bold=True)
            continue
        if text.startswith("关键词："):
            content = text.split("：", 1)[1]
            clear_paragraph(paragraph)
            label_run = paragraph.add_run("关键词：")
            value_run = paragraph.add_run(content)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.5
            set_run_text_font(label_run, "宋体", Pt(12), bold=False)
            set_run_text_font(value_run, "宋体", Pt(12), bold=False)
            continue
        if text.startswith("Keywords:"):
            content = text[len("Keywords:"):].strip()
            clear_paragraph(paragraph)
            label_run = paragraph.add_run("Keywords:")
            spacer_run = paragraph.add_run(" ")
            value_run = paragraph.add_run(content)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.5
            set_run_text_font(label_run, "Times New Roman", Pt(12), bold=True)
            set_run_text_font(spacer_run, "Times New Roman", Pt(12), bold=False)
            set_run_text_font(value_run, "Times New Roman", Pt(12), bold=False)
            continue
        if re.match(r"^图\s*\d+\.\d+", text):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = Pt(20)
            for run in paragraph.runs:
                set_run_text_font(run, "黑体", Pt(12), bold=False)
            continue
        if text.startswith("图片来源："):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = Pt(18)
            for run in paragraph.runs:
                set_run_text_font(run, "宋体", Pt(10.5), bold=False)
            continue
        if re.match(r"^表\s*\d+\.\d+", text):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = Pt(20)
            for run in paragraph.runs:
                set_run_text_font(run, "黑体", Pt(12), bold=False)
            continue
        if text.startswith("Sindustry") or text.startswith("Sleader"):
            equation_text = "Sindustry=0.35×Zm+0.35×Zf+0.15×Zv-0.15×Zr"
            equation_number = "（4.2.1）"
            if text.startswith("Sleader"):
                equation_text = "Sleader=100×(0.20×s1+0.15×s2+0.25×s3+0.20×s4+0.10×s5+0.10×s6)"
                equation_number = "（4.4.1）"
            clear_paragraph(paragraph)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = BODY_LINE_SPACING
            usable_width = doc.sections[-1].page_width - doc.sections[-1].left_margin - doc.sections[-1].right_margin
            paragraph.paragraph_format.tab_stops.add_tab_stop(Emu(int(usable_width / 2)), alignment=WD_TAB_ALIGNMENT.CENTER)
            paragraph.paragraph_format.tab_stops.add_tab_stop(Emu(int(usable_width)), alignment=WD_TAB_ALIGNMENT.RIGHT)
            run = paragraph.add_run(f"\t{equation_text}\t{equation_number}")
            set_run_text_font(run, "Times New Roman", Pt(10.5), bold=False)
            continue
        if re.match(r"^\[\d+\]", text):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = Cm(-1.27)
            paragraph.paragraph_format.left_indent = Cm(1.27)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = BODY_LINE_SPACING
            for run in paragraph.runs:
                set_run_text_font(run, "宋体", Pt(12), bold=False)
            continue
        if text.startswith("题   目：") or text.startswith("学    院：") or text.startswith("专    业："):
            continue

        paragraph.style = "Normal"
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.first_line_indent = BODY_FIRST_LINE_INDENT
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        if abstract_index < idx < english_abstract_index or english_abstract_index < idx < toc_index:
            paragraph.paragraph_format.line_spacing = ABSTRACT_LINE_SPACING
        else:
            paragraph.paragraph_format.line_spacing = BODY_LINE_SPACING
        for run in paragraph.runs:
            if english_abstract_index < idx < toc_index:
                set_run_text_font(run, "Times New Roman", Pt(12), bold=False)
            else:
                format_body_run(run)
        apply_superscript_citations(paragraph)


def configure_headers_and_footers(doc: Document) -> None:
    for index, section in enumerate(doc.sections):
        if index < 2:
            configure_section_header(section, None)
            configure_section_footer(section, False)
            continue
        configure_section_header(section, HEADER_TEXT)
        configure_section_footer(section, True)
    if len(doc.sections) >= 3:
        set_section_page_number_format(doc.sections[2], fmt="upperRoman", start=1)
    if len(doc.sections) >= 4:
        set_section_page_number_format(doc.sections[3], fmt="decimal", start=1)


def compute_toc_pages_from_pdf(pdf_path: Path) -> dict[str, str]:
    reader = PdfReader(str(pdf_path))
    raw_page_texts = [page.extract_text() or "" for page in reader.pages]
    page_texts = [normalize_search_text(text) for text in raw_page_texts]
    page_lines = [
        [normalize_search_text(line) for line in text.splitlines() if line.strip()]
        for text in raw_page_texts
    ]

    toc_start_index = next(
        (i for i, text in enumerate(page_texts) if text.count(".") > 100 and "1绪论" in text),
        None,
    )
    chapter_start_index = None
    if toc_start_index is not None:
        chapter_start_index = next(
            (
                i
                for i in range(toc_start_index + 1, len(page_texts))
                if page_texts[i].count(".") < 50 and "1绪论" in page_texts[i]
            ),
            None,
        )
    if chapter_start_index is None:
        chapter_start_index = next(
            (i for i, text in enumerate(page_texts) if text.count(".") < 50 and "1绪论" in text),
            None,
        )
    if chapter_start_index is None:
        raise RuntimeError("Failed to determine the first body page for TOC rebuilding.")

    toc_pages = {
        "摘  要": "I",
        "ABSTRACT": "II",
    }
    for _, title in TOC_BLUEPRINT:
        if title in toc_pages:
            continue
        normalized_title = normalize_search_text(title)
        for page_index in range(chapter_start_index, len(page_texts)):
            if normalized_title in page_lines[page_index]:
                toc_pages[title] = str(page_index - chapter_start_index + 1)
                break
    return toc_pages


def remove_page_break_paragraph_before(doc: Document, title: str) -> None:
    paragraphs = doc.paragraphs
    for idx, paragraph in enumerate(paragraphs):
        if paragraph.text.strip() != title:
            continue
        if idx == 0:
            return
        previous = paragraphs[idx - 1]
        if previous.text.strip() == "" and 'w:type="page"' in previous._element.xml:
            delete_paragraph(previous)
        paragraph.paragraph_format.page_break_before = True
        return


def normalize_major_breaks(doc: Document) -> None:
    heading_1_paragraphs = [paragraph for paragraph in doc.paragraphs if paragraph.style.name == "Heading 1"]
    for index, paragraph in enumerate(heading_1_paragraphs):
        previous = paragraph._element.getprevious()
        if previous is not None and previous.tag == qn("w:p"):
            from docx.text.paragraph import Paragraph

            previous_para = Paragraph(previous, paragraph._parent)
            if previous_para.text.strip() == "" and 'w:type="page"' in previous_para._element.xml:
                delete_paragraph(previous_para)
        if index == 0:
            paragraph.paragraph_format.page_break_before = False
            remove_page_break_before_flag(paragraph)
        else:
            paragraph.paragraph_format.page_break_before = True

    for title in ["结 论", "参考文献", "致 谢"]:
        remove_page_break_paragraph_before(doc, title)


def harmonize_section_breaks(target_doc: Document, source_doc: Document) -> None:
    source_chapter_break = deepcopy(find_heading_break_paragraph(source_doc, "绪论")._element.pPr.find(qn("w:sectPr")))
    target_chapter_break_paragraph = find_heading_break_paragraph(target_doc, "绪论")
    target_chapter_break_ppr = target_chapter_break_paragraph._element.get_or_add_pPr()
    existing_target_break = target_chapter_break_ppr.find(qn("w:sectPr"))
    if existing_target_break is not None:
        target_chapter_break_ppr.remove(existing_target_break)
    target_chapter_break_ppr.append(source_chapter_break)

    heading_index = None
    for index, paragraph in enumerate(target_doc.paragraphs):
        if paragraph.text.strip() == "绪论":
            heading_index = index
            break
    if heading_index is None:
        raise RuntimeError("Failed to locate chapter heading for section cleanup.")

    for index in range(max(0, heading_index - 4), heading_index):
        paragraph = target_doc.paragraphs[index]
        if paragraph._element is target_chapter_break_paragraph._element:
            continue
        p_pr = paragraph._element.pPr
        if p_pr is None:
            continue
        sect_pr = p_pr.find(qn("w:sectPr"))
        if sect_pr is not None:
            p_pr.remove(sect_pr)

    target_body = target_doc._element.body
    target_body_sect = target_body.find(qn("w:sectPr"))
    if target_body_sect is not None:
        target_body.remove(target_body_sect)
    target_body.append(deepcopy(source_doc._element.body.find(qn("w:sectPr"))))


def compose_official_template(source_doc: Document, source_layout_doc: Document) -> Document:
    TMP_DOC_DIR.mkdir(parents=True, exist_ok=True)

    front_doc = Document(str(TEMPLATE_PATH))
    fill_cover(front_doc)
    fill_declaration_table(front_doc)
    remove_elements_from_paragraph(front_doc, "摘  要")

    remove_elements_before_paragraph(source_doc, "摘  要")

    composer = Composer(front_doc)
    composer.append(source_doc)
    composer.save(str(COMPOSED_TMP_PATH))

    composed_doc = Document(str(COMPOSED_TMP_PATH))
    harmonize_section_breaks(composed_doc, source_layout_doc)
    set_core_properties(composed_doc)
    return composed_doc


def export_submission_artifacts(doc_path: Path) -> Path:
    template_page_1, template_page_2 = ensure_template_pdf_pages()
    front_cover_png = TMP_FRONT_ASSET_DIR / "front_cover.png"
    declaration_png = TMP_FRONT_ASSET_DIR / "front_declaration.png"
    front_pdf = TMP_FRONT_ASSET_DIR / "front_pages.pdf"

    cover_render = render_cover_page(template_page_1, front_cover_png)
    declaration_render = render_declaration_page(template_page_2, declaration_png)
    build_front_pdf(front_cover_png, declaration_png, front_pdf)

    body_pdf = export_docx_pdf(doc_path, TMP_DOCX_RENDER_DIR)
    merge_submission_pdf(front_pdf, body_pdf, OUTPUT_PDF_PATH)
    if LEGACY_DUPLICATE_PDF_PATH.exists():
        LEGACY_DUPLICATE_PDF_PATH.unlink()
    render_pdf_preview(OUTPUT_PDF_PATH, TMP_SUBMISSION_RENDER_DIR)
    validate_front_page_alignment(
        get_rendered_preview_page(TMP_SUBMISSION_RENDER_DIR, 1),
        cover_render["placements"],
        cover_render["image_size"],
    )
    validate_declaration_page_alignment(
        get_rendered_preview_page(TMP_SUBMISSION_RENDER_DIR, 2),
        declaration_render["placements"],
        declaration_render["label_boxes"],
        declaration_render["image_size"],
    )
    return OUTPUT_PDF_PATH


def main() -> None:
    source_layout_doc = Document(str(DOC_PATH))
    normalize_chapter_headings(source_layout_doc)
    update_body_content(source_layout_doc)
    replace_formula_tables(source_layout_doc)
    relayout_figures(source_layout_doc)
    normalize_major_breaks(source_layout_doc)

    composed_doc = compose_official_template(source_layout_doc, Document(str(DOC_PATH)))
    normalize_reference_entries(composed_doc)
    format_data_tables(composed_doc)
    configure_headers_and_footers(composed_doc)
    format_paragraphs(composed_doc)
    apply_minor_layout_overrides(composed_doc)
    composed_doc.save(str(DOC_PATH))

    pdf_path = export_submission_artifacts(DOC_PATH)
    toc_pages = compute_toc_pages_from_pdf(pdf_path)
    rebuild_toc(composed_doc, toc_pages)
    format_paragraphs(composed_doc)
    apply_minor_layout_overrides(composed_doc)
    composed_doc.save(str(DOC_PATH))

    pdf_path = export_submission_artifacts(DOC_PATH)
    verified_toc_pages = compute_toc_pages_from_pdf(pdf_path)
    if verified_toc_pages != toc_pages:
        rebuild_toc(composed_doc, verified_toc_pages)
        format_paragraphs(composed_doc)
        apply_minor_layout_overrides(composed_doc)
        composed_doc.save(str(DOC_PATH))
        export_submission_artifacts(DOC_PATH)


if __name__ == "__main__":
    main()
