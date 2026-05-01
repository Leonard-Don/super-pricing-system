from __future__ import annotations

import shutil
import subprocess
from pathlib import Path
from urllib.request import Request, urlopen

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from PIL import Image


PROJECT_ROOT = Path(__file__).resolve().parents[1]
DOC_OUTPUT_DIR = PROJECT_ROOT / "output" / "doc"
TMP_DOC_DIR = PROJECT_ROOT / "tmp" / "docs"

OUTPUT_DOCX = DOC_OUTPUT_DIR / "上海大学本科毕业论文_基于大数据的热门行业识别与龙头股遴选研究_通信学院附录材料.docx"
OUTPUT_PDF = OUTPUT_DOCX.with_suffix(".pdf")

TMP_RENDER_DIR = TMP_DOC_DIR / "comm_appendices_render"
SOURCE_CACHE_DIR = TMP_DOC_DIR / "appendix_source"
SOURCE_PDF = SOURCE_CACHE_DIR / "oecd_ai_ml_big_data_finance.pdf"
SOURCE_PAGE_DIR = TMP_RENDER_DIR / "source_pages"
FLATTEN_PAGE_DIR = TMP_RENDER_DIR / "flatten_pages"

INTERMEDIATE_PDF = TMP_RENDER_DIR / OUTPUT_PDF.name

SOURCE_URL = (
    "https://www.oecd.org/content/dam/oecd/en/publications/reports/2021/08/"
    "artificial-intelligence-machine-learning-and-big-data-in-finance_8d088cbb/98e761e7-en.pdf"
)
SOURCE_TITLE_EN = "Artificial Intelligence, Machine Learning and Big Data in Finance"
SOURCE_TITLE_CN = "OECD《Artificial Intelligence, Machine Learning and Big Data in Finance》节选译文"
SOURCE_NOTE = (
    "原文来源：OECD（2021）《Artificial Intelligence, Machine Learning and Big Data in Finance》。"
    "本附录选取报告第7、15、21、37页作为英文原稿节选，按学院模板要求以清晰截图形式插入。"
)
SOURCE_PAGE_SPECS = [
    (8, 7),
    (16, 15),
    (22, 21),
    (38, 37),
]

STUDENT_NAME = "唐梓涵"
COMPLETION_PLACE = "上海"
COMPLETION_DATE = "2026年4月17日"

APPENDIX_A_TRANSLATION_PARAGRAPHS = [
    (
        "执行摘要首先指出，金融领域中的人工智能并不是一个抽象概念，而是能够围绕人类预设目标进行预测、"
        "推荐或决策的机器系统。随着另类数据来源不断增加，金融机构开始把大量数据分析能力与机器学习模型结合起来，"
        "使系统能够在经验和数据的基础上自动提升预测效果，而不再完全依赖人工显式编程。报告认为，新冠疫情进一步"
        "加速了这种数字化趋势。全球人工智能支出预计将在2020年至2024年间显著增长，而资产管理、算法交易、"
        "信贷承做以及基于区块链的金融服务，正是在数据愈加丰富、算力更加可得且成本下降的背景下迅速推进的。"
    ),
    (
        "报告同时强调，人工智能在金融业扩张的意义不仅在于“用了新技术”，更在于它改变了金融机构获取竞争优势的方式。"
        "一方面，人工智能能够通过改善决策流程、自动化执行、提升风险管理和合规效率、优化后台流程等方式降低成本、"
        "提高生产率，并最终改善机构盈利能力；另一方面，它又能通过产品和服务的定制化、个性化以及新产品供给来提升"
        "金融服务质量。由此带来的竞争优势并不仅仅有利于机构本身，也可能通过更好的产品质量、更丰富的选择和更低的"
        "服务成本传导给金融消费者。"
    ),
    (
        "但 OECD 在同一页中就提醒，政策制定者不能只看到效率收益，而忽视人工智能在金融中的风险外溢。"
        "人工智能应用可能制造或放大金融风险与非金融风险，并引出消费者保护和投资者保护层面的新问题。"
        "由于不少人工智能模型缺乏充分的可解释性和可理解性，监管者和机构内部治理框架可能难以及时弄清模型为何得出"
        "某个结果，这不仅会影响单个机构的稳健运行，也可能在市场层面放大顺周期性和系统性风险。报告还指出，"
        "人工智能可能带来带有偏见、不公平甚至歧视性的消费者结果，并伴随数据管理与使用层面的担忧。"
    ),
    (
        "在第一章导论部分，报告进一步把人工智能、机器学习和大数据的关系讲得更清楚。其核心意思是："
        "人工智能系统依靠海量数据源和数据分析能力运行，而机器学习模型则从这些数据集中学习，使模型能够在不被人工"
        "逐项编程的情况下不断“自我改进”。随着这一技术体系渗透到医疗、汽车、消费品和物联网等多个产业，"
        "金融服务提供者也正在更广泛地把它部署到零售银行、公司银行、资产管理、交易、保险、监管科技和监管者科技中。"
        "例如，银行可以用它提供定制化产品、客户服务、授信评分和欺诈监测；资产管理机构可以用它做投顾、组合管理和"
        "风险管理；保险机构则可以用于理赔管理。"
    ),
    (
        "这一章还强调，随着基于大数据的人工智能和机器学习在金融服务中的重要性不断提高，由此产生的风险也更值得"
        "政策制定者持续关注。报告提到，许多国家和国际组织已经开始讨论，监管者和监督者应如何确保人工智能应用在"
        "金融领域的风险处于可接受范围内。也就是说，人工智能在金融中的推广不能只靠市场自发推进，而需要与监管讨论、"
        "治理结构、模型验证和责任界定同步展开。这一点说明，技术能力越强，制度和治理安排反而越不能缺席。"
    ),
    (
        "在第二章关于“人工智能、机器学习和大数据对金融业务模式与活动影响”的讨论中，OECD 认为，"
        "金融业采用人工智能的直接驱动力来自两个方面：一是金融服务内部可利用的数据规模不断扩大，二是人工智能和"
        "机器学习为机构带来的预期竞争优势越来越明确。海量数据分析能力与更便宜的计算资源，尤其是云计算，"
        "使机器学习模型能够识别数据中的信号并捕捉隐藏关系，而这种能力往往超出了人类在传统分析框架下的处理上限。"
        "因此，人工智能不只是辅助报表阅读的工具，而是在重塑资产管理、投资、交易、放贷以及区块链金融等多类活动的"
        "业务组织方式。"
    ),
    (
        "报告配图中列举的应用场景也说明了这种重塑的广度。前台环节可以利用人工智能进行资产配置、智能投顾、"
        "客户服务、生物识别认证、交易策略执行和个性化产品推荐；中后台则可能将其用于反洗钱、了解你的客户、"
        "信用评分、风险承做、数据分析、对账、报送、记录管理和控制流程。换言之，人工智能并不是局限于某一个高频交易"
        "算法，而是逐步嵌入金融机构从客户触达、风控审批到后台运营的完整链条之中。正因为它渗透得如此广泛，"
        "其收益和风险都具有跨部门、跨流程的连锁效应。"
    ),
    (
        "到了第三章，报告把关注点转向风险与缓释。文中明确提出，随着人工智能、机器学习和大数据技术在金融市场中的"
        "应用规模和应用谱系不断扩大，越来越多的挑战与风险正在被识别出来，这些问题可能发生在数据层面、模型与机构层面，"
        "也可能发生在社会层面和系统层面。报告逐项点出了值得重点关注的议题，包括数据管理与集中度问题、偏见与歧视风险、"
        "可解释性、模型稳健性和韧性、治理与问责、监管适配问题，以及就业和技能结构变化等。也就是说，"
        "人工智能在金融中越成功，其治理问题就越不应被视为附属问题。"
    ),
    (
        "在数据管理部分，报告把数据放在所有人工智能应用的核心位置来讨论。它承认，大数据和机器学习确实能够提高效率、"
        "降低成本，并通过更高质量的产品和服务提高客户满意度，但同时也指出，数据质量、数据隐私与保密、网络安全以及"
        "公平性问题，会使这些金融产品和服务暴露于重要的非金融风险。尤其是在授信等场景中，如果模型使用了不恰当的数据，"
        "或者对某些群体代表性不足，就可能无意中造成偏见与歧视。报告还特别强调，训练、测试与验证数据对于模型是否能够"
        "在尾部事件情形下保持预测能力至关重要；如果数据只覆盖平稳时期，模型在压力状态下的可靠性就会明显下降。"
    ),
    (
        "报告还进一步提醒，大数据在金融中的广泛使用可能改变数据控制权和竞争格局。如果高价值数据、算力资源和模型训练能力"
        "过度集中于少数大型平台或机构，中小机构在开发、验证和部署人工智能应用时就会处于不利位置，由此带来竞争与集中度方面"
        "的新问题。与此同时，消费者往往难以准确理解自己的数据在多大范围内被收集、加工、共享和再次利用，一旦数据授权边界、"
        "隐私保护安排或跨机构共享机制不够清晰，服务便利性的提升就可能以透明度下降和信任损耗为代价。"
    ),
    (
        "此外，报告并没有把模型风险简单理解为“算法黑箱”，而是把它直接和治理责任联系起来。模型结果如果不能被业务部门、"
        "风险管理部门以及监管者适度解释，就很难在授信、投资组合管理、交易执行等关键流程中建立清晰的问责链条。"
        "人工智能系统越深入金融核心业务，人类监督、独立复核和异常纠偏机制就越不能缺位，否则一旦模型出现系统性偏差，"
        "机构不仅难以及时修正结果，也难以明确责任归属。"
    ),
    (
        "基于这些观察，报告实际上给出了一个十分明确的政策含义：金融业在利用人工智能、机器学习和大数据获取效率收益时，"
        "必须同步强化数据治理、模型验证、内部责任划分和外部监管沟通。对行业参与者而言，真正重要的不是简单地“引入模型”，"
        "而是确保所使用的数据真实、适当、可审计，确保模型结果在业务流程中可被解释、可被复核，并且在消费者保护、"
        "市场公平和机构稳健性之间保持平衡。只有当技术部署与治理约束共同推进时，人工智能在金融中的应用才能持续释放价值，"
        "而不是在效率提升的同时积累更大的隐性风险。"
    ),
]

APPENDIX_B_SECTIONS = {
    "一、课题目的和意义": [
        (
            "本课题围绕“基于大数据的热门行业识别与龙头股遴选研究”展开，其直接目的在于把金融数据处理、"
            "量化分析方法和前后端软件实现结合起来，形成一个能够持续运行、可解释且便于展示的研究原型系统。"
            "从学术层面看，课题帮助学生理解多源金融数据的组织方式、行业横截面比较方法以及龙头股多维评价逻辑，"
            "将原本分散的大数据、机器学习、可视化和软件工程知识整合到同一个具体问题中。"
        ),
        (
            "从工程层面看，本课题的意义在于证明一个本科毕业设计不必停留在静态报告或单一算法演示上，而可以"
            "围绕真实场景构建具备数据获取、评分分析、接口服务和前端展示能力的完整系统。对于金融研究工作而言，"
            "热门行业和龙头企业的识别本身具有较高的现实价值，因为它关系到研究效率、信息组织能力以及不同类型"
            "市场参与者对行业变化的理解方式。通过多源数据适配和结构化展示，课题尝试降低研究人员在重复检索与"
            "手工整理中的时间成本。"
        ),
        (
            "从社会意义看，本课题并不仅仅服务于狭义的投资判断。行业热度识别系统如果设计得足够透明，实际上也"
            "可以作为金融教育、行业观察和信息整理工具，帮助使用者更系统地认识产业结构变化、资金关注方向和不同"
            "指标之间的关系。它有助于推动数据素养和理性分析意识，而不是单纯强化“追热点”的投机思维。因此，"
            "课题的意义既包括技术与学术价值，也包括提高信息理解效率、促进理性决策和服务数字化研究转型的社会价值。"
        ),
    ],
    "二、课题的作用与思考": [
        (
            "从积极作用看，本课题的第一项价值在于提升研究效率并增强结果透明度。传统行业分析往往依赖人工翻阅"
            "行情终端、研报和表格，难以同时兼顾速度、覆盖范围与解释能力。系统把价格动量、资金流、交易活跃度、"
            "估值和财务指标等关键信息组织到统一界面后，研究者能够在较短时间内完成横截面对比，并通过评分拆解了解"
            "某一行业或个股被识别出来的原因。这种透明化过程对教学、训练和研究协同都具有积极作用。"
        ),
        (
            "第二项积极作用在于，多源数据适配和历史快照记录为后续复盘提供了基础。相比一次性的静态分析，系统"
            "能够保留行业热力图快照、观察配置与详情结果，这有助于研究者在后续回看市场结构时发现不同时间窗口下"
            "行业表现的差异，从而提升研究工作的连续性和可验证性。对于学校培养目标而言，这种基于真实工程环境的"
            "训练更接近未来数字金融、金融科技和数据分析岗位所要求的综合能力。"
        ),
        (
            "但与此同时，本课题也带来值得认真思考的潜在问题。首先，任何评分模型都存在指标选择与权重设定偏差，"
            "如果使用者过度依赖系统输出，而忽略宏观环境、政策变化和企业基本面细节，就可能把“辅助研究工具”误解"
            "为“自动决策机器”。其次，多源数据虽然提高了可用性，但不同接口之间仍可能存在口径不一致、字段缺失或"
            "刷新延迟，若系统没有清楚提示数据来源和局限性，结果就可能被错误解读。"
        ),
        (
            "此外，从社会影响角度看，类似系统如果被大规模用于零门槛的投资推荐，可能强化部分投资者的从众行为，"
            "促使用户只关注榜单排序而忽视风险教育。再从可持续发展角度看，持续抓取、缓存和计算金融数据虽然不属于"
            "高能耗产业，但仍然意味着网络资源、算力和存储资源的消耗。如果系统未来规模扩大，就需要进一步考虑资源"
            "利用效率、接口合规性和数据治理问题。因此，本课题的价值并不只是把系统做出来，更重要的是在设计与使用"
            "过程中始终保持技术审慎和社会责任意识。"
        ),
    ],
    "三、课题的建议与构想": [
        (
            "针对上述影响，首先建议在系统层面继续强化可解释性与风险提示功能。除了展示综合得分外，还应继续保留"
            "维度拆解、数据来源标签、历史快照对比和字段完整性说明，使使用者能够区分“数据充分支持的结果”和“依赖"
            "代理变量或补充源得到的结果”。如果未来扩展到更多用户场景，还可以加入更明确的研究提示，例如提醒用户"
            "评分结果仅用于研究参考，不构成直接投资建议。"
        ),
        (
            "其次，建议后续研究把行业识别结果与社会环境和可持续发展视角结合得更紧密。目前系统主要围绕价格、资金、"
            "估值和财务指标展开，未来可以进一步纳入政策文本、产业链关系、环境责任信息和企业长期治理指标，以避免"
            "系统只强调短期热度而忽视产业质量。对于通信、半导体、新能源等高景气行业，这种改进尤其有必要，因为它"
            "能帮助研究者同时看到市场热度与长期发展质量之间的关系。"
        ),
        (
            "再次，建议在工程实现上持续优化数据治理与资源利用方式。系统应进一步完善数据源校验、异常日志记录、"
            "缓存更新策略和字段一致性检查，减少因数据漂移造成的误判。同时，在可行条件下采用更合理的调度频率、"
            "增量更新方式和轻量化快照存储策略，以降低无效请求与重复计算带来的资源消耗。这样既能提升系统稳定性，"
            "也符合面向长期运行的可持续设计思路。"
        ),
        (
            "最后，从人才培养角度出发，建议把本课题继续扩展为“研究工具 + 教学案例 + 风险教育”三位一体的实践平台。"
            "它不仅可以服务毕业设计展示，也可以用于帮助学生理解金融数据分析、行业比较和软件系统设计之间的联系。"
            "如果后续能在课程教学、实验平台和研究训练中进一步沉淀标准化案例，本课题的价值就不再局限于一次性的毕业"
            "成果，而能转化为更持久的教学与研究资源。"
        ),
    ],
}


def set_east_asia_font(run, name: str) -> None:
    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    r_pr.rFonts.set(qn("w:eastAsia"), name)


def count_chinese_chars(text: str) -> int:
    return sum(1 for ch in text if "\u4e00" <= ch <= "\u9fff")


def validate_content_requirements() -> None:
    appendix_a_count = count_chinese_chars("".join(APPENDIX_A_TRANSLATION_PARAGRAPHS))
    appendix_b_count = count_chinese_chars(
        "".join("".join(paragraphs) for paragraphs in APPENDIX_B_SECTIONS.values())
    )
    if appendix_a_count < 2500:
        raise ValueError(f"附录A中文译文不足2500字，当前仅 {appendix_a_count} 字。")
    if appendix_b_count < 1000:
        raise ValueError(f"附录B课题调研报告不足1000字，当前仅 {appendix_b_count} 字。")


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr_text)
    run._r.append(fld_end)
    set_east_asia_font(run, "Times New Roman")
    run.font.size = Pt(10.5)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.6)
    section.bottom_margin = Cm(2.4)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)

    header = section.header.paragraphs[0]
    header.text = "上海大学本科毕业论文（设计）附录材料"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if header.runs:
        set_east_asia_font(header.runs[0], "宋体")
        header.runs[0].font.size = Pt(10.5)

    footer = section.footer.paragraphs[0]
    add_page_number(footer)


def style_body(paragraph, first_line_indent: bool = True, align=WD_ALIGN_PARAGRAPH.JUSTIFY) -> None:
    paragraph.alignment = align
    fmt = paragraph.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.first_line_indent = Pt(21) if first_line_indent else Pt(0)
    for run in paragraph.runs:
        set_east_asia_font(run, "宋体")
        run.font.size = Pt(12)


def add_body_paragraph(doc: Document, text: str, first_line_indent: bool = True) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    set_east_asia_font(run, "宋体")
    run.font.size = Pt(12)
    style_body(paragraph, first_line_indent=first_line_indent)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    set_east_asia_font(run, "黑体")
    run.font.bold = True
    run.font.size = Pt(18 if level == 1 else 16 if level == 2 else 14)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(8)
    fmt.space_after = Pt(6)
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    fmt.first_line_indent = Pt(0)


def add_center_title(doc: Document, text: str, english: bool = False) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    set_east_asia_font(run, "Times New Roman" if english else "黑体")
    run.font.bold = True
    run.font.size = Pt(14 if english else 16)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(6)
    fmt.space_after = Pt(6)
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE


def add_caption(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    set_east_asia_font(run, "宋体")
    run.font.size = Pt(10.5)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE


def add_image_page(doc: Document, image_path: Path, page_label: int) -> None:
    add_caption(doc, f"英文原稿节选，第{page_label}页")
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Cm(15.0))


def add_page_break(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def add_signature_block(doc: Document) -> None:
    for line in [
        f"作者：{STUDENT_NAME}",
        f"完成地点：{COMPLETION_PLACE}",
        COMPLETION_DATE,
    ]:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(line)
        set_east_asia_font(run, "宋体")
        run.font.size = Pt(12)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)


def ensure_source_pdf() -> None:
    SOURCE_CACHE_DIR.mkdir(parents=True, exist_ok=True)
    if SOURCE_PDF.exists() and SOURCE_PDF.stat().st_size > 100_000:
        return
    request = Request(
        SOURCE_URL,
        headers={"User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7)"},
    )
    with urlopen(request) as response, SOURCE_PDF.open("wb") as file_obj:
        shutil.copyfileobj(response, file_obj)


def render_source_pages() -> list[tuple[int, Path]]:
    ensure_source_pdf()
    SOURCE_PAGE_DIR.mkdir(parents=True, exist_ok=True)
    rendered_pages: list[tuple[int, Path]] = []
    for page_index, page_label in SOURCE_PAGE_SPECS:
        prefix = SOURCE_PAGE_DIR / f"page_{page_label}"
        for old_file in SOURCE_PAGE_DIR.glob(f"{prefix.name}-*.png"):
            old_file.unlink()
        subprocess.run(
            [
                "pdftoppm",
                "-r",
                "170",
                "-png",
                "-f",
                str(page_index + 1),
                "-l",
                str(page_index + 1),
                str(SOURCE_PDF),
                str(prefix),
            ],
            check=True,
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
        )
        matches = sorted(SOURCE_PAGE_DIR.glob(f"{prefix.name}-*.png"))
        if not matches:
            raise FileNotFoundError(f"未成功渲染英文原稿第{page_label}页。")
        rendered_pages.append((page_label, matches[0]))
    return rendered_pages


def build_appendix_doc(rendered_source_pages: list[tuple[int, Path]]) -> Document:
    doc = Document()
    configure_document(doc)

    add_heading(doc, "附  录", level=1)
    add_body_paragraph(
        doc,
        "以下材料依据上海大学通信学院本科毕业设计撰写模板整理，包含附录A“英译汉”和附录B“课题调研报告”两部分内容。",
    )

    add_page_break(doc)

    add_heading(doc, "附录A  英译汉", level=1)
    add_body_paragraph(doc, "一、英文原文", first_line_indent=False)
    add_center_title(doc, SOURCE_TITLE_EN, english=True)
    add_body_paragraph(doc, SOURCE_NOTE, first_line_indent=False)
    add_page_break(doc)

    for index, (page_label, image_path) in enumerate(rendered_source_pages):
        add_image_page(doc, image_path, page_label)
        if index != len(rendered_source_pages) - 1:
            add_page_break(doc)

    add_page_break(doc)

    add_body_paragraph(doc, "二、英文翻译", first_line_indent=False)
    add_center_title(doc, SOURCE_TITLE_CN)
    for paragraph in APPENDIX_A_TRANSLATION_PARAGRAPHS:
        add_body_paragraph(doc, paragraph)

    add_page_break(doc)

    add_heading(doc, "附录B  课题调研报告", level=1)
    for section_title, paragraphs in APPENDIX_B_SECTIONS.items():
        add_body_paragraph(doc, section_title, first_line_indent=False)
        for paragraph in paragraphs:
            add_body_paragraph(doc, paragraph)

    add_signature_block(doc)
    return doc


def export_pdf(docx_path: Path, pdf_path: Path) -> None:
    pdf_path.parent.mkdir(parents=True, exist_ok=True)
    subprocess.run(
        [
            "soffice",
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(pdf_path.parent),
            str(docx_path),
        ],
        check=True,
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
    )


def flatten_pdf_to_image_pdf(source_pdf: Path, output_pdf: Path) -> None:
    FLATTEN_PAGE_DIR.mkdir(parents=True, exist_ok=True)
    for old_file in FLATTEN_PAGE_DIR.glob("*"):
        old_file.unlink()
    prefix = FLATTEN_PAGE_DIR / "appendix_page"
    subprocess.run(
        ["pdftoppm", "-r", "170", "-png", str(source_pdf), str(prefix)],
        check=True,
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
    )
    page_images = sorted(FLATTEN_PAGE_DIR.glob("appendix_page-*.png"))
    if not page_images:
        raise FileNotFoundError("未能生成用于贴图版PDF的页面图像。")

    rgb_images = []
    for page_image in page_images:
        with Image.open(page_image) as image:
            rgb_images.append(image.convert("RGB"))

    first_image, *rest_images = rgb_images
    output_pdf.parent.mkdir(parents=True, exist_ok=True)
    first_image.save(output_pdf, save_all=True, append_images=rest_images, resolution=170.0)
    for image in rgb_images:
        image.close()


def main() -> None:
    validate_content_requirements()

    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    if TMP_RENDER_DIR.exists():
        shutil.rmtree(TMP_RENDER_DIR)
    TMP_RENDER_DIR.mkdir(parents=True, exist_ok=True)

    rendered_source_pages = render_source_pages()
    doc = build_appendix_doc(rendered_source_pages)
    doc.save(str(OUTPUT_DOCX))

    export_pdf(OUTPUT_DOCX, INTERMEDIATE_PDF)
    flatten_pdf_to_image_pdf(INTERMEDIATE_PDF, OUTPUT_PDF)


if __name__ == "__main__":
    main()
